"""
Test Card Viewer Component with DOCX Export
Generate and display test cards directly from uploaded documents in vector database.
Enhanced with Word document export capability.
"""

import streamlit as st
import pandas as pd
import json
from config.settings import config
from lib.api.client import api_client
from services.chromadb_service import chromadb_service
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


def create_test_card_docx(test_cards_data, metadata, doc_title):
    """
    Create a Word document with test card tables.

    Args:
        test_cards_data: Dictionary of section_title -> card_content
        metadata: Test card metadata
        doc_title: Document title

    Returns:
        BytesIO object containing the DOCX file
    """
    doc = Document()

    # Add title
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Sections: {metadata.get('total_sections', 0)}")
    doc.add_paragraph(f"Total Tests: {metadata.get('total_tests', 0)}")
    doc.add_paragraph("")

    # Add table of contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("(Update TOC in Word: Right-click → Update Field)")
    doc.add_page_break()

    # Process each section
    for section_title, card_content in test_cards_data.items():
        # Section heading
        doc.add_heading(section_title, level=1)

        # Try to parse as JSON for structured data
        try:
            test_cards = json.loads(card_content)

            if test_cards and isinstance(test_cards, list):
                # Create table with header row
                table = doc.add_table(rows=1, cols=10)
                table.style = 'Light Grid Accent 1'

                # Header row
                header_cells = table.rows[0].cells
                headers = ['Test ID', 'Test Title', 'Procedures', 'Expected Results',
                          'Acceptance Criteria', 'Dependencies', 'Executed', 'Pass', 'Fail', 'Notes']

                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    # Set header cell background
                    header_cells[i]._element.get_or_add_tcPr().append(
                        header_cells[i]._element.new_tag('w:shd', w_fill='4472C4')
                    )

                # Add test card rows
                for tc in test_cards:
                    row_cells = table.add_row().cells

                    # Test ID
                    row_cells[0].text = tc.get('test_id', 'N/A')

                    # Test Title
                    row_cells[1].text = tc.get('test_title', 'N/A')

                    # Procedures (numbered list)
                    procedures = tc.get('procedures', [])
                    row_cells[2].text = '\n'.join([f"{i+1}. {p}" for i, p in enumerate(procedures)])

                    # Expected Results
                    row_cells[3].text = tc.get('expected_results', 'N/A')

                    # Acceptance Criteria
                    row_cells[4].text = tc.get('acceptance_criteria', 'N/A')

                    # Dependencies
                    deps = tc.get('dependencies', [])
                    row_cells[5].text = ', '.join(deps) if deps else 'None'

                    # Execution tracking columns (checkboxes)
                    row_cells[6].text = '☐'
                    row_cells[7].text = '☐'
                    row_cells[8].text = '☐'
                    row_cells[9].text = ''

                    # Set font size for all cells in row
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)

                # Add spacing after table
                doc.add_paragraph("")

        except (json.JSONDecodeError, TypeError):
            # If not JSON or parse fails, treat as markdown
            if card_content and len(card_content.strip()) > 50:
                # Add markdown content as paragraph
                doc.add_paragraph(card_content)
            else:
                doc.add_paragraph("No test cards found for this section.")

        # Page break between sections
        doc.add_page_break()

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return bio


def TestCardViewer():
    """
    Component for viewing and managing test cards generated from uploaded documents.
    Creates test cards directly from specifications/standards in the vector database.
    """
    st.header("Test Card Generator")
    st.info("Generate executable test cards from uploaded specification documents (MIL-STD, ISO, etc.)")

    # Get available collections
    if "collections" not in st.session_state:
        st.session_state.collections = chromadb_service.get_collections()

    collections = st.session_state.collections

    if not collections:
        st.warning("No collections found. Please upload documents first in the Upload Documents page.")
        st.stop()

    # Document selection
    st.subheader("Select Source Documents")

    # Collection selection
    selected_collection = st.selectbox(
        "Select Collection (containing specifications/standards):",
        collections,
        key="testcard_collection"
    )

    # Load documents from collection
    if st.button("Load Documents", key="load_docs"):
        with st.spinner(f"Loading documents from {selected_collection}..."):
            try:
                documents = chromadb_service.get_documents(selected_collection)
                st.session_state.available_docs = [
                    {
                        'document_id': doc.document_id,
                        'document_name': doc.document_name,
                        'file_type': doc.file_type,
                        'total_chunks': doc.total_chunks
                    }
                    for doc in documents
                ]
                st.success(f"Loaded {len(documents)} documents")
            except Exception as e:
                st.error(f"Failed to load documents: {e}")
                st.session_state.available_docs = []

    # Document selection
    available_docs = st.session_state.get("available_docs", [])
    if available_docs:
        doc_options = {
            f"{doc['document_name']} ({doc['file_type']}) - {doc['total_chunks']} chunks": doc['document_id']
            for doc in available_docs
        }

        selected_docs = st.multiselect(
            "Select Document(s) to Generate Test Cards From:",
            list(doc_options.keys()),
            key="testcard_selected_docs"
        )

        selected_doc_ids = [doc_options[name] for name in selected_docs]

        if selected_doc_ids:
            st.success(f"Selected {len(selected_doc_ids)} document(s)")
        else:
            st.info("Select at least one document to generate test cards")
    else:
        st.info("Click 'Load Documents' to see available documents")
        st.stop()

    if not selected_doc_ids:
        st.warning("Please select at least one document to continue")
        st.stop()

    # Generate test plan first (which creates sections)
    st.markdown("---")
    st.subheader("Generate Test Plan with Test Cards")

    col1, col2 = st.columns(2)
    with col1:
        test_card_format = st.radio(
            "Test Card Format:",
            ["markdown_table", "json"],
            horizontal=True,
            key="testcard_format"
        )
    with col2:
        doc_title = st.text_input(
            "Test Plan Title:",
            value="Test Plan",
            key="testcard_doc_title"
        )

    generate_btn = st.button("Generate Test Plan & Cards", type="primary", key="generate_test_cards")

    if generate_btn:
        with st.spinner("Step 1/2: Generating test plan from documents... This may take several minutes."):
            try:
                # Step 1: Generate test plan using multi-agent service
                test_plan_response = api_client.post(
                    f"{config.endpoints.doc_gen}/generate_documents",
                    data={
                        "source_collections": [selected_collection],
                        "source_doc_ids": selected_doc_ids,
                        "doc_title": doc_title,
                        "use_rag": True,
                        "top_k": 5
                    },
                    # timeout=600  # 10 minutes for test plan generation
                )

                documents = test_plan_response.get("documents", [])
                if not documents:
                    st.error("No test plan generated. Please check your documents.")
                    st.stop()

                st.success("✅ Step 1/2: Test plan generated successfully!")

            except Exception as e:
                st.error(f"❌ Failed to generate test plan: {e}")
                import traceback
                st.code(traceback.format_exc())
                st.stop()

        with st.spinner("Step 2/2: Generating test cards from sections..."):
            try:
                # Step 2: Generate test cards from the document content
                test_cards = {}

                for doc in documents:
                    section_title = doc.get('title', 'Document Section')
                    content = doc.get('content', '')

                    if content:
                        # Call generate-test-card API for each section
                        card_response = api_client.post(
                            f"{config.endpoints.doc_gen}/generate-test-card",
                            data={
                                "section_title": section_title,
                                "rules_markdown": content,
                                "format": test_card_format
                            },
                            # timeout=180
                        )

                        test_cards[section_title] = card_response.get("test_card_content", "")

                st.session_state.test_cards = test_cards
                st.session_state.test_card_metadata = {
                    "total_sections": len(test_cards),
                    "total_tests": sum(
                        len([line for line in content.split('\n') if line.strip().startswith('| TC-')])
                        for content in test_cards.values()
                    ) if test_card_format == "markdown_table" else 0,
                    "format": test_card_format,
                    "source_collection": selected_collection,
                    "source_docs": selected_doc_ids,
                    "doc_title": doc_title
                }

                st.success(f"✅ Step 2/2: Generated {len(test_cards)} test card sections!")

            except Exception as e:
                st.error(f"❌ Failed to generate test cards: {e}")
                import traceback
                st.code(traceback.format_exc())
                st.session_state.test_cards = {}

    # Display test cards
    if st.session_state.get("test_cards"):
        st.markdown("---")
        st.subheader("Test Cards")

        metadata = st.session_state.get("test_card_metadata", {})

        # Summary metrics
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Sections", metadata.get("total_sections", 0))
        with col2:
            st.metric("Total Tests", metadata.get("total_tests", 0))
        with col3:
            st.metric("Format", metadata.get("format", "N/A").upper())

        st.markdown("---")

        # Display each section's test cards
        for section_title, card_content in st.session_state.test_cards.items():
            with st.expander(f"{section_title}", expanded=False):

                # Display based on format
                if metadata.get("format") == "json":
                    try:
                        test_cards = json.loads(card_content)

                        if test_cards:
                            # Create interactive display for each test card
                            for idx, tc in enumerate(test_cards):
                                st.markdown(f"**{tc.get('test_id', f'TC-{idx+1}')}: {tc.get('test_title', 'Untitled Test')}**")

                                col1, col2 = st.columns([2, 1])

                                with col1:
                                    st.markdown("**Procedures:**")
                                    for i, proc in enumerate(tc.get('procedures', []), 1):
                                        st.markdown(f"{i}. {proc}")

                                with col2:
                                    st.markdown("**Expected Results:**")
                                    st.write(tc.get('expected_results', 'N/A'))

                                    st.markdown("**Acceptance Criteria:**")
                                    st.write(tc.get('acceptance_criteria', 'N/A'))

                                # Dependencies
                                if tc.get('dependencies'):
                                    st.markdown(f"**Dependencies:** {', '.join(tc['dependencies'])}")

                                # Execution tracking
                                st.markdown("**Test Execution:**")
                                exec_col1, exec_col2, exec_col3, exec_col4 = st.columns(4)

                                with exec_col1:
                                    st.checkbox("Executed", key=f"exec_{section_title}_{idx}")
                                with exec_col2:
                                    st.checkbox("Pass", key=f"pass_{section_title}_{idx}")
                                with exec_col3:
                                    st.checkbox("Fail", key=f"fail_{section_title}_{idx}")
                                with exec_col4:
                                    st.text_input("Notes", key=f"notes_{section_title}_{idx}", label_visibility="collapsed", placeholder="Notes...")

                                st.markdown("---")

                            # Section export options
                            st.markdown("### Export Options")
                            export_col1, export_col2 = st.columns(2)

                            with export_col1:
                                # Convert to DataFrame for CSV export
                                df = pd.DataFrame(test_cards)
                                csv_data = df.to_csv(index=False).encode('utf-8')
                                st.download_button(
                                    label="Download as CSV",
                                    data=csv_data,
                                    file_name=f"{section_title.replace(' ', '_')}_test_cards.csv",
                                    mime="text/csv",
                                    key=f"csv_{section_title}"
                                )

                            with export_col2:
                                # Download as JSON
                                json_data = json.dumps(test_cards, indent=2).encode('utf-8')
                                st.download_button(
                                    label="Download as JSON",
                                    data=json_data,
                                    file_name=f"{section_title.replace(' ', '_')}_test_cards.json",
                                    mime="application/json",
                                    key=f"json_{section_title}"
                                )

                        else:
                            st.info("No test cards found for this section")

                    except json.JSONDecodeError as e:
                        st.error(f"Failed to parse JSON test cards: {e}")
                        st.code(card_content, language="json")

                else:  # markdown_table format
                    if card_content and len(card_content.strip()) > 50:
                        st.markdown(card_content)

                        # Download markdown
                        st.download_button(
                            label="Download as Markdown",
                            data=card_content.encode('utf-8'),
                            file_name=f"{section_title.replace(' ', '_')}_test_cards.md",
                            mime="text/markdown",
                            key=f"md_{section_title}"
                        )
                    else:
                        st.info("No test cards found for this section")

        # Global export options
        st.markdown("---")
        st.subheader("Export All Test Cards")

        export_format = st.radio(
            "Export Format:",
            ["Word (DOCX)", "Excel (CSV)"],
            horizontal=True,
            key="global_export_format"
        )

        if st.button("Export All", type="primary", key="export_all_button"):
            with st.spinner(f"Generating {export_format} export..."):
                try:
                    if export_format == "Word (DOCX)":
                        # Generate DOCX with test card tables
                        bio = create_test_card_docx(
                            st.session_state.test_cards,
                            metadata,
                            metadata.get('doc_title', doc_title)
                        )

                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        st.download_button(
                            label="Download Word Document",
                            data=bio,
                            file_name=f"test_cards_{timestamp}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx_all"
                        )
                        st.success(f"✅ Word document ready: test_cards_{timestamp}.docx")

                    else:  # Excel (CSV)
                        # Combine all test cards into single CSV
                        all_cards = []
                        for section_title, card_content in st.session_state.test_cards.items():
                            try:
                                if metadata.get("format") == "json":
                                    cards = json.loads(card_content)
                                    for card in cards:
                                        card['section'] = section_title
                                    all_cards.extend(cards)
                            except:
                                pass

                        if all_cards:
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            df_all = pd.DataFrame(all_cards)
                            csv_data = df_all.to_csv(index=False).encode('utf-8')

                            st.download_button(
                                label="Download CSV",
                                data=csv_data,
                                file_name=f"test_cards_{timestamp}.csv",
                                mime="text/csv",
                                key="download_csv_all"
                            )
                            st.success(f"✅ CSV export ready: {len(all_cards)} test cards")
                        else:
                            st.warning("No test cards available to export")

                except Exception as e:
                    st.error(f"❌ Export failed: {e}")
                    import traceback
                    st.code(traceback.format_exc())

    else:
        st.info("👆 Click 'Generate Test Plan & Cards' to generate test cards from the selected documents")

    # Help section
    with st.expander("Help & Information"):
        st.markdown("""
        ### How to Use Test Card Generator

        1. **Select Source Documents**:
           - Choose a collection from your vector database
           - Load and select specification documents (MIL-STD, ISO, etc.)

        2. **Generate Test Plan & Cards**:
           - Choose output format (Markdown table or JSON)
           - Provide a test plan title
           - Click to generate - this creates a test plan and extracts test cards

        3. **Review Test Cards**:
           - Expand sections to view individual test cards with:
             - Test ID and Title
             - Step-by-step procedures
             - Expected results and acceptance criteria
             - Dependencies (if any)

        4. **Track Execution**:
           - Use checkboxes to mark tests as executed, passed, or failed
           - Add notes for each test

        5. **Export**:
           - Download individual sections (CSV, JSON, Markdown)
           - Export all test cards together (Word DOCX, Excel CSV)

        ### Export Formats

        - **Word (DOCX)**: Professional Word document with formatted tables, headers, and TOC
        - **Excel (CSV)**: Spreadsheet format for test tracking in Excel
        - **JSON**: Machine-readable format for integration with test management tools
        - **Markdown**: For manual conversion or further processing

        ### Tips

        - **JSON format** provides the most interactive display with execution tracking
        - **Word (DOCX)** format creates professional test documentation with proper formatting
        - Use **CSV** format for test execution tracking in spreadsheet applications
        - Test cards are generated directly from your uploaded specification documents
        """)
