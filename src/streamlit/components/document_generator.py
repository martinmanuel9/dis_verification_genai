import streamlit as st
import requests
from components.upload_documents import *
from utils import *


FASTAPI_API = os.getenv("FASTAPI_URL", "http://localhost:9020")

def Document_Generator():
    st.header("Document Generator")
    st.info("Upload templates, and generate comprehensive documentation using AI analysis.")
    
    # Document Generator sub-modes
    doc_gen_mode = st.radio(
        "Select Action:",
        ["Rule Development Agents", "Template Management", "Generate Documents", "Generated Documents"],
        horizontal=True
    )
    
    # ----------------------------------------------------------------------
    # RULE DEVELOPMENT AGENTS SUB-MODE
    # ----------------------------------------------------------------------
    if doc_gen_mode == "Rule Development Agents":
        st.subheader("Rule Development Agents")
        st.info("Create specialized agents for extracting rules, requirements, and test plans from technical documents.")
        
        # Enhanced rule development templates
        rule_agent_templates = {
            "Rule Extraction Agent": {
                "description": "Specialized agent for extracting detailed, testable rules and requirements from technical documents",
                "system_prompt": """You are a test planning expert specializing in extracting comprehensive, testable rules from technical documents.

                Your expertise includes:
                1. **Rule Identification**: Identify EVERY possible testable requirement which usually contains "shall", "must", "may", "will", "could" or "should"
                2. **Detailed Analysis**: Extract rules that are extremely detailed, explicit, and step-by-step
                3. **Measurable Criteria**: Include specific measurements, acceptable ranges, and referenced figures/tables
                4. **Test Strategy**: For ambiguous requirements, describe specific test strategies
                5. **Dependency Analysis**: Identify dependencies between rules and requirements
                6. **Conflict Detection**: Detect and resolve conflicts between requirements

                **Output Format Requirements:**
                - Use markdown headings and bolded text for organization
                - Use the provided template to structure the output
                - Include all relevant details from the analysis

                **Analysis Approach:**
                - Extract both explicit and implicit requirements
                - Consider edge cases and boundary conditions
                - Identify verification and validation methods
                - Note any missing information that would affect testing
                - Provide specific test procedures where applicable""",

                "user_prompt": """Develop a test plan based on the template provided and use the following document:

                {data_sample}

                to develop a comprehensive test plan for each requirement with a verification method, verification approach, necessary test steps.""",

                "temperature": 0.2,
                "max_tokens": 2500
            },
            
            "Test Plan Synthesis Agent": {
                "description": "Agent specialized in combining multiple rule sets into comprehensive test plans",
                "system_prompt": """You are a senior QA documentation engineer specializing in synthesizing complex test plans from multiple rule sources.

                Your expertise includes:
                1. **Test Plan Integration**: Combine multiple rule sets into coherent test plans
                2. **Cross-Reference Analysis**: Identify overlapping content and merge similar steps
                3. **Dependency Management**: Map dependencies between different sections and requirements
                4. **Conflict Resolution**: Identify and resolve conflicts between different rule sources
                5. **Test Organization**: Structure tests in logical execution order
                6. **Coverage Analysis**: Ensure comprehensive coverage of all requirements

                **Synthesis Methodology:**
                - Merge similar test steps and eliminate redundancy
                - Cross-reference overlapping content between sections
                - Organize tests by logical execution sequence
                - Identify prerequisite tests and setup requirements
                - Group related test procedures for efficiency
                - Maintain traceability to original requirements

                **Output Standards:**
                - Use content-based titles that reflect actual test scope
                - Maintain markdown formatting with clear sections
                - Provide explicit step-by-step test procedures
                - Include setup, execution, and verification steps
                - Note any special equipment or conditions required
                - Cross-reference related test procedures""",
                            
                            "user_prompt": """You are provided with detailed test rules from multiple sections. Synthesize these into a single, comprehensive test plan:

                {data_sample}

                Create a combined test plan with this structure:

                ## [Content-Based Test Plan Title]

                **Test Dependencies:**
                - List prerequisite tests and setup requirements
                - Note any equipment or environmental conditions needed

                **Conflict Resolution:**
                - Address any conflicts between different rule sources
                - Provide recommended resolution approaches

                **Integrated Test Procedures:**
                1. [Comprehensive, step-by-step test procedures]
                2. [Merge similar steps, eliminate redundancy]
                3. [Organize in logical execution order]
                4. [Include setup, execution, and verification phases]

                **Cross-References:**
                - Map relationships between different test procedures
                - Note shared requirements and common verification steps

                Focus on creating a test plan that is:
                - Logically organized and executable
                - Comprehensive in coverage
                - Efficient in execution order
                - Clear in requirements and procedures""",
            
                "temperature": 0.3,
                "max_tokens": 3000
            },
            
            "Document Section Analyzer": {
                "description": "Agent for analyzing document sections and preparing structured analysis",
                "system_prompt": """You are a technical document analysis expert specializing in structured content extraction and preparation.

                Your capabilities include:
                1. **Content Classification**: Identify types of content (requirements, procedures, specifications, etc.)
                2. **Section Analysis**: Extract key topics, themes, and technical focus areas
                3. **Structure Mapping**: Understand document hierarchy and relationships
                4. **Content Preparation**: Prepare content for further analysis by other specialized agents
                5. **Metadata Extraction**: Identify references, figures, tables, and cross-references

                **Analysis Framework:**
                - Identify the primary purpose and scope of each section
                - Extract technical specifications and requirements
                - Note procedural steps and methodologies
                - Identify measurement criteria and acceptance standards
                - Map relationships to other document sections
                - Highlight areas needing further clarification

                **Content Organization:**
                - Categorize content by type (functional, performance, interface, etc.)
                - Identify compliance requirements and standards references
                - Extract numerical values, ranges, and specifications
                - Note any conditional or situational requirements
                - Highlight critical vs. optional requirements""",
                            
                            "user_prompt": """Analyze the following document section and provide structured analysis:

                {data_sample}

                Provide analysis in this format:

                ## Section Analysis: [Content-Based Title]

                **Content Type:**
                - Identify the primary type of content (requirements, procedures, specifications, etc.)

                **Key Topics:**
                - List main topics and technical focus areas
                - Note any specialized terminology or concepts

                **Technical Specifications:**
                - Extract specific measurements, values, and criteria
                - List any referenced standards or specifications

                **Requirements Identified:**
                - Functional requirements
                - Performance requirements  
                - Interface requirements
                - Constraint requirements

                **References and Dependencies:**
                - Note any figures, tables, or cross-references mentioned
                - Identify dependencies on other sections or documents

                **Analysis Notes:**
                - Areas requiring clarification
                - Potential ambiguities or interpretation issues
                - Recommendations for further analysis

                This analysis will be used by specialized rule extraction agents.""",
                
                "temperature": 0.2,
                "max_tokens": 2000
            }
        }
        
        available_models = get_available_models_cached()
        
        # Manual Agent Creation Section - WITH PROPER TEMPLATE HANDLING
        st.subheader("Agent Creation")
        st.info("Create individual specialized agents with custom configurations.")
        
        # Template selection OUTSIDE the form to allow dynamic updates
        col1_pre, col2_pre = st.columns([2, 1])
        
        with col1_pre:
            selected_template = st.selectbox(
                "Choose Agent Template:",
                ["Custom"] + list(rule_agent_templates.keys()),
                key="template_selector",
                help="Select a template to auto-populate the form fields"
            )
        
        with col2_pre:
            if selected_template != "Custom":
                template_info = rule_agent_templates[selected_template]
                st.info(f"**{selected_template}**: {template_info['description']}")
        
        # Get template defaults based on selection
        if selected_template != "Custom":
            template = rule_agent_templates[selected_template]
            default_system_prompt = template["system_prompt"]
            default_user_prompt = template["user_prompt"]
            default_temp = template.get("temperature", 0.3)
            default_tokens = template.get("max_tokens", 2000)
            default_name = f"Custom {selected_template}"
        else:
            default_system_prompt = ""
            default_user_prompt = ""
            default_temp = 0.3
            default_tokens = 2000
            default_name = ""
        
        # Form with proper template integration
        with st.form("create_custom_rule_agent", clear_on_submit=False):
            col1, col2 = st.columns([1, 1])
            
            with col1:
                rule_agent_name = st.text_input(
                    "Agent Name",
                    value=default_name,
                    placeholder="e.g., Custom Standards Extractor"
                )
                
                # Model selection
                if available_models:
                    rule_agent_model = st.selectbox("Model", available_models, key="custom_model")
                else:
                    st.error("No models available")
                    rule_agent_model = None
                
                # Advanced settings in the left column
                st.subheader("Configuration")
                
                rule_temperature = st.slider(
                    "Temperature", 
                    0.0, 1.0, 
                    default_temp, 
                    0.1,
                    help="Lower = more consistent, Higher = more creative"
                )
                
                rule_max_tokens = st.number_input(
                    "Max Tokens", 
                    100, 4000, 
                    default_tokens, 
                    100,
                    help="Maximum response length"
                )
            
            with col2:
                rule_system_prompt = st.text_area(
                    "System Prompt (Define Agent's Role & Expertise)",
                    value=default_system_prompt,
                    height=300,
                    help="Define the agent's expertise and approach",
                    placeholder="You are an expert specializing in..."
                )
                
                rule_user_prompt = st.text_area(
                    "User Prompt Template",
                    value=default_user_prompt,
                    height=200,
                    help="Must include {data_sample} placeholder",
                    placeholder="Analyze the following:\n\n{data_sample}\n\nProvide analysis..."
                )
            
            # Validation indicators
            col1_val, col2_val = st.columns(2)
            
            with col1_val:
                if rule_agent_name and len(rule_agent_name.strip()) >= 3:
                    st.success("Agent name valid")
                else:
                    st.error("Agent name too short")
            
            with col2_val:
                if "{data_sample}" in rule_user_prompt:
                    st.success("User prompt has {data_sample}")
                else:
                    st.error("Missing {data_sample} placeholder")
            
            # Create button
            col1_btn, col2_btn, col3_btn = st.columns([1, 2, 1])
            
            with col2_btn:
                rule_create_submitted = st.form_submit_button(
                    "Create Custom Agent", 
                    type="primary",
                    use_container_width=True
                )
            
            if rule_create_submitted:
                # Comprehensive validation
                validation_errors = []
                
                if not rule_agent_name or len(rule_agent_name.strip()) < 3:
                    validation_errors.append("Agent name must be at least 3 characters")
                
                if not rule_system_prompt or len(rule_system_prompt.strip()) < 50:
                    validation_errors.append("System prompt must be substantial (at least 50 characters)")
                
                if "{data_sample}" not in rule_user_prompt:
                    validation_errors.append("User prompt must include {data_sample} placeholder")
                
                if not rule_agent_model:
                    validation_errors.append("Please select a valid model")
                
                if validation_errors:
                    st.error("**Validation Errors:**")
                    for error in validation_errors:
                        st.error(f"• {error}")
                else:
                    payload = {
                        "name": rule_agent_name.strip(),
                        "model_name": rule_agent_model,
                        "system_prompt": rule_system_prompt.strip(),
                        "user_prompt_template": rule_user_prompt.strip(),
                        "temperature": rule_temperature,
                        "max_tokens": rule_max_tokens
                    }
                    
                    try:
                        with st.spinner("Creating custom rule development agent..."):
                            response = requests.post(f"{FASTAPI_API}/create-agent", json=payload, timeout=30)
                            
                            if response.status_code == 200:
                                result = response.json()
                                st.success(f"Custom agent '{rule_agent_name}' created successfully!")
                                st.info(f"Agent ID: {result.get('agent_id', 'Unknown')}")
                                st.info(f"Model: {rule_agent_model}")
                                st.info(f"Temperature: {rule_temperature}")
                                
                                # Clear cached data to show new agent
                                if 'current_rule_agents' in st.session_state:
                                    del st.session_state.current_rule_agents
                                if 'agents_data' in st.session_state:
                                    st.session_state.agents_data = []
                                
                            else:
                                error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                
                                if "already exists" in error_detail.lower():
                                    st.error(f"Agent name '{rule_agent_name}' already exists. Please choose a different name.")
                                else:
                                    st.error(f"Failed to create agent: {error_detail}")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        
        st.markdown("---")
        
        # Agent Management Section
        st.subheader("Manage Rule Development Agents")
        
        col1_mgmt, col2_mgmt = st.columns([1, 1])
        
        with col1_mgmt:
            if st.button("Load All Agents", key="load_all_agents"):
                try:
                    with st.spinner("Loading all agents..."):
                        agents_response = requests.get(f"{FASTAPI_API}/get-agents", timeout=10)
                        if agents_response.status_code == 200:
                            all_agents = agents_response.json().get("agents", [])
                            st.session_state.current_rule_agents = all_agents
                            st.success(f"Loaded {len(all_agents)} total agents")
                        else:
                            st.error("Failed to load agents")
                except Exception as e:
                    st.error(f"Error: {e}")
        
        with col2_mgmt:
            if 'current_rule_agents' in st.session_state:
                total_agents = len(st.session_state.current_rule_agents)
                rule_related = sum(1 for agent in st.session_state.current_rule_agents 
                                if any(keyword in agent['name'].lower() 
                                    for keyword in ['rule', 'test', 'document', 'analysis', 'extract', 'auto']))
                st.metric("Total Agents", total_agents, delta=f"{rule_related} rule-related")
        
        # Display current agents
        if 'current_rule_agents' in st.session_state and st.session_state.current_rule_agents:
            agents = st.session_state.current_rule_agents
            
            # Filter for rule development related agents
            rule_agents = [agent for agent in agents 
            if any(keyword in agent['name'].lower() for keyword in ['rule', 'test', 'document', 'analysis', 'extract', 'synthesis', 'auto'])]
            
            if rule_agents:
                st.subheader(f"Rule Development Agents ({len(rule_agents)} found)")
                
                # Create enhanced display table
                agent_display_data = []
                for agent in rule_agents:
                    agent_display_data.append({
                        "Name": agent.get("name", "Unknown"),
                        "Model": agent.get("model_name", "Unknown"),
                        "Temperature": f"{agent.get('temperature', 0.7):.1f}",
                        "Max Tokens": agent.get("max_tokens", 1000),
                        "Queries": agent.get("total_queries", 0),
                        "Status": "Active" if agent.get("is_active", True) else "Inactive",
                        "Created": agent.get("created_at", "Unknown")[:10] if agent.get("created_at") else "Unknown"
                    })
                
                st.dataframe(agent_display_data, use_container_width=True, height=300)
                
                # Quick Test Section
                st.subheader("Test Agent")
                
                if rule_agents:
                    test_agent_choice = st.selectbox(
                        "Select agent to test:",
                        [f"{agent['name']} (ID: {agent['id']})" for agent in rule_agents],
                        key="test_agent_select"
                    )
                    
                    # Sample test content
                    test_content = st.text_area(
                        "Test Content:",
                        value="""4.2.3 Signal Processing Requirements
                        The system SHALL process incoming RF signals according to the following specifications:
                        4.2.3.1 Frequency Range: The system SHALL operate within the frequency range of 30 MHz to 3 GHz with a tolerance of ±0.1%.
                        4.2.3.2 Signal Sensitivity: The minimum detectable signal level SHALL be -110 dBm or better across the entire frequency range.
                        4.2.3.3 Processing Time: Signal processing SHALL be completed within 50 milliseconds from signal acquisition to output generation.""",
                        height=200,
                        help="Edit this content to test your agent"
                    )
                    
                    if st.button("Run Test", type="secondary"):
                        if test_agent_choice and test_content:
                            # Extract agent ID from selection
                            agent_id = int(test_agent_choice.split("ID: ")[1].rstrip(")"))
                            agent_name = test_agent_choice.split(" (ID:")[0]
                            
                            with st.spinner(f"Testing {agent_name}..."):
                                try:
                                    payload = {
                                        "data_sample": test_content,
                                        "agent_ids": [agent_id]
                                    }
                                    
                                    response = requests.post(
                                        f"{FASTAPI_API}/compliance-check",
                                        json=payload,
                                        timeout=60
                                    )
                                    
                                    if response.status_code == 200:
                                        result = response.json()
                                        st.success("Test completed successfully!")
                                        
                                        # Show test results
                                        details = result.get("details", {})
                                        for idx, analysis in details.items():
                                            st.subheader(f"Results from {analysis.get('agent_name', 'Unknown Agent')}")
                                            st.markdown(analysis.get("reason", "No analysis generated"))
                                            
                                            if "response_time_ms" in result:
                                                st.caption(f"Response time: {result['response_time_ms']}ms")
                                    else:
                                        st.error(f"Test failed: HTTP {response.status_code}")
                                        error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                        st.error(f"Error details: {error_detail}")
                                        
                                except Exception as e:
                                    st.error(f"Test error: {str(e)}")
            
            else:
                st.info("No rule development agents found. Create some using the options above!")
        
        else:
            st.info("Click 'Load All Agents' to see existing agents and their status.")
        
        # Footer
        st.markdown("---")
        st.info("**Next Steps**: After creating agents, use them in 'Generate Documents' for document analysis!")
    
    
    # ----------------------------------------------------------------------
    # TEMPLATE MANAGEMENT SUB-MODE  
    # ----------------------------------------------------------------------
    
    if doc_gen_mode == "Template Management":
        st.header("Document Template Management")
        st.info("Upload and manage document templates for automated rule generation and test plan creation.")
        
        collections = get_chromadb_collections()

        render_upload_component(
            available_collections= collections,
            load_collections_func= get_chromadb_collections,
            create_collection_func= create_collection,
            upload_endpoint=f"{CHROMADB_API}/documents/upload-and-process",
            job_status_endpoint=f"{CHROMADB_API}/jobs/{{job_id}}"
        )
        
    # ----------------------------------------------------------------------
    # GENERATE DOCUMENTS SUB-MODE
    # ----------------------------------------------------------------------
    elif doc_gen_mode == "Generate Documents":
        st.subheader("Generate Test Plans")

        # ——————————————————————————
        # 1) Pick agents
        # ——————————————————————————
        agents = st.session_state.get("available_rule_agents") or requests.get(f"{FASTAPI_API}/get-agents").json()["agents"]
        # rule_agents = [a for a in agents if "rule" in a["name"].lower()]
        agent_map = {f"{a['name']} ({a['model_name']})": a["id"] for a in agents}
        selected_agents = st.multiselect("Select Agents", list(agent_map.keys()), key="gen_agents")
        if not selected_agents:
            st.info("Choose at least one agent to proceed"); st.stop()

        # ——————————————————————————
        # 2a) Pick TEMPLATE collection & load template docs
        # ——————————————————————————
        template_collection = st.selectbox(
            "Template Collection (the one you uploaded as templates)",
            st.session_state.collections,
            key="gen_template_coll",
        )
        if st.button("Load Template Library", key="gen_load_templates"):
            st.session_state.template_docs = get_all_documents_in_collection(template_collection)

        template_docs = st.session_state.get("template_docs", [])
        template_map = {d["document_name"]: d["document_id"] for d in template_docs}
        selected_templates = st.multiselect(
            "Select Template(s)", list(template_map.keys()), key="gen_templates"
        )
        template_doc_ids = [template_map[name] for name in selected_templates]

        # ——————————————————————————
        # 2b) Pick SOURCE collection & load source docs
        # ——————————————————————————
        source_collection = st.selectbox(
            "Source Collection (your requirements/standards)",
            st.session_state.collections,
            key="gen_source_coll",
        )
        if st.button("Load Source Documents", key="gen_load_sources"):
            st.session_state.source_docs = get_all_documents_in_collection(source_collection)

        source_docs = st.session_state.get("source_docs", [])
        source_map = {d["document_name"]: d["document_id"] for d in source_docs}
        selected_sources = st.multiselect(
            "Select Source Document(s)", list(source_map.keys()), key="gen_sources"
        )
        source_doc_ids = [source_map[name] for name in selected_sources]

        # ——————————————————————————
        # 3) Pick agents 
        # ——————————————————————————
        agent_ids = [agent_map[label] for label in selected_agents]

        # ——————————————————————————
        # 4) Let user name the output file
        # ——————————————————————————
        out_name = st.text_input(
            "Output file name (no extension):",
            value="Generated_Test_Plan",
            key="gen_filename"
        ).strip()

        # ——————————————————————————
        # 5) Generate analyses
        # ——————————————————————————
        if st.button("Generate Documents", type="primary"):
            if not template_doc_ids or not source_doc_ids:
                st.error("You must select at least one templated and one source doc.")
            else:
                payload = {
                    "template_collection": template_collection,
                    "template_doc_ids":    template_doc_ids,
                    "source_collections":  [source_collection],
                    "source_doc_ids":      source_doc_ids,
                    "agent_ids":           agent_ids,
                    "use_rag":             True,
                    "top_k":               5,
                    "doc_title":           out_name
                }
                st.write("about to call /generate_documents on", FASTAPI_API)
                st.write("Payload:", payload)
                with st.spinner("Calling Document Generator…"):
                    try:
                        resp = requests.post(
                            f"{FASTAPI_API}/generate_documents",
                            json=payload
                            # timeout=300
                        )
                        # now resp is guaranteed to exist
                        if not resp.ok:
                            st.error(f"Error {resp.status_code}: {resp.text}")
                        else:
                            docs = resp.json().get("documents", [])
                            st.success(f"Generated {len(docs)} documents")
                            for d in docs:
                                blob = base64.b64decode(d["docx_b64"])
                                st.download_button(
                                    label=f"{d['title']}.docx",
                                    data=blob,
                                    file_name=f"{d['title']}.docx",
                                    mime=(
                                        "application/"
                                        "vnd.openxmlformats-"
                                        "officedocument."
                                        "wordprocessingml.document"
                                    )
                                )
                    except Exception as e:
                            st.error("Request exception: " + str(e))

        # ——————————————————————————
        # 6) Offer download once we have results
        # ——————————————————————————
        if st.session_state.get("gen_results"):
            buffer = build_docx_bytes([
                {
                    "document_title": r["title"],
                    "analysis_content": r["content"],
                    "source_document": "",
                    "agent_name": ""
                }
                for r in st.session_state.gen_results
            ])
            st.download_button(
                "Download Combined DOCX",
                data=buffer.getvalue(),
                file_name=f"{out_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


    # ----------------------------------------------------------------------
    # Sub GENERATED DOCUMENTS SUBMODE
    # ----------------------------------------------------------------------

    elif doc_gen_mode == "Generated Documents":
        st.subheader("Generated Test Plans")
        st.info("View, manage, and export your generated rule analysis documents.")
        
        # Check if there are generated results in session state
        if 'generated_results' in st.session_state and st.session_state.generated_results:
            results = st.session_state.generated_results
            output_collection = st.session_state.get('output_collection', 'generated-documents')
            
            st.success(f"Found {len(results)} generated documents from your last generation session")
            
            # Results Overview
            st.subheader("Generation Overview")
            
            col1_overview, col2_overview, col3_overview, col4_overview = st.columns(4)
            
            with col1_overview:
                st.metric("Total Documents", len(results))
            
            with col2_overview:
                unique_agents = len(set(r['agent_name'] for r in results))
                st.metric("Agents Used", unique_agents)
            
            with col3_overview:
                unique_sources = len(set(r['source_document'] for r in results))
                st.metric("Source Documents", unique_sources)
            
            with col4_overview:
                total_chars = sum(r['content_length'] for r in results)
                st.metric("Total Content", f"{total_chars:,} chars")
            
            # Detailed Results Table
            st.subheader("Generated Documents")
            
            # Create results table
            results_table = []
            for i, result in enumerate(results):
                results_table.append({
                    "Index": i + 1,
                    "Document Title": result['document_title'],
                    "Source Document": result['source_document'],
                    "Agent": result['agent_name'],
                    "Content Length": f"{result['content_length']:,} chars",
                    "Generated": result['generation_timestamp'][:19],
                    "Processing Time": f"{result.get('processing_time_ms', 0)}ms"
                })
            
            st.dataframe(results_table, use_container_width=True, height=400)
            
            # Document Actions
            st.markdown("---")
            st.subheader("Document Actions")
            
            # Document selector for individual actions
            doc_titles = [r['document_title'] for r in results]
            selected_doc_title = st.selectbox(
                "Select document for actions:",
                ["--Select Document--"] + doc_titles,
                key="selected_generated_doc"
            )
            
            if selected_doc_title != "--Select Document--":
                # Find the selected result
                selected_result = next((r for r in results if r['document_title'] == selected_doc_title), None)
                
                if selected_result:
                    col1_action, col2_action, col3_action = st.columns(3)
                    
                    with col1_action:
                        if st.button("Preview Document", key="preview_generated_doc"):
                            st.subheader(f"Preview: {selected_result['document_title']}")
                            
                            # Document metadata
                            with st.expander("Document Information", expanded=False):
                                st.json({
                                    "Source Document": selected_result['source_document'],
                                    "Agent": selected_result['agent_name'],
                                    "Generated": selected_result['generation_timestamp'],
                                    "Content Length": f"{selected_result['content_length']:,} characters",
                                    "Processing Time": f"{selected_result.get('processing_time_ms', 0)}ms"
                                })
                            
                            # Content preview
                            content = selected_result['analysis_content']
                            st.text_area(
                                "Generated Content:",
                                content,
                                height=400,
                                disabled=True,
                                key="preview_content_area"
                            )
                    
                    with col2_action:
                        if st.button("Download as DOCX", key="download_single_docx"):
                            try:
                                # Create DOCX document
                                doc = Document()
                                doc.add_heading(selected_result['document_title'], 0)
                                
                                # Add metadata section
                                doc.add_heading('Document Information', 1)
                                doc.add_paragraph(f"Source Document: {selected_result['source_document']}")
                                doc.add_paragraph(f"Generated by Agent: {selected_result['agent_name']}")
                                doc.add_paragraph(f"Generated on: {selected_result['generation_timestamp']}")
                                doc.add_paragraph(f"Content Length: {selected_result['content_length']:,} characters")
                                
                                # Add main content
                                doc.add_heading('Analysis Content', 1)
                                
                                content = selected_result['analysis_content']
                                markdown_to_docx(content, doc)
                                
                                # Save to temporary file
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                                doc.save(temp_file.name)
                                
                                # Create download
                                with open(temp_file.name, "rb") as file:
                                    st.download_button(
                                        label="Download DOCX",
                                        data=file.read(),
                                        file_name=f"{selected_result['document_title'].replace(' ', '_')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_single_docx_btn"
                                    )
                                
                                # Cleanup
                                os.unlink(temp_file.name)
                                st.success("Download ready!")
                                
                            except Exception as e:
                                st.error(f"Download error: {e}")
                    
                    with col3_action:
                        if st.button("Save to Collection", key="save_to_collection"):
                            try:
                                # Create a temporary file with the content
                                temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding='utf-8')
                                temp_file.write(selected_result['analysis_content'])
                                temp_file.close()
                                
                                # Read the file back as bytes for upload
                                with open(temp_file.name, 'rb') as file:
                                    file_data = file.read()
                                
                                # Create a file-like object for the upload
                                from io import BytesIO
                                file_obj = BytesIO(file_data)
                                file_obj.name = f"{selected_result['document_title']}.txt"
                                
                                # Store in ChromaDB
                                result = store_files_in_chromadb(
                                    [file_obj], 
                                    output_collection,
                                    model_name="basic",
                                    chunk_size=2000,
                                    chunk_overlap=200,
                                    store_images=False
                                )
                                
                                st.success(f"Document saved to collection '{output_collection}'!")
                                st.json(result)
                                
                                # Cleanup
                                os.unlink(temp_file.name)
                                
                            except Exception as e:
                                st.error(f"Save error: {e}")
            
            st.markdown("---")
            
            # Bulk Actions
            st.subheader("Bulk Actions")
            
            col1_bulk, col2_bulk, col3_bulk = st.columns(3)
            
            with col1_bulk:
                if st.button("Download All as Combined DOCX", key="download_all_combined"):
                    try:
                        from docx import Document
                        import tempfile
                        import os
                        
                        # Create combined document
                        doc = Document()
                        doc.add_heading('Generated Rule Analysis Documents', 0)
                        
                        # Add generation summary
                        doc.add_heading('Generation Summary', 1)
                        doc.add_paragraph(f"Total Documents: {len(results)}")
                        doc.add_paragraph(f"Agents Used: {unique_agents}")
                        doc.add_paragraph(f"Source Documents: {unique_sources}")
                        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        doc.add_page_break()
                        
                        # Add each document
                        for i, result in enumerate(results, 1):
                            doc.add_heading(f"{i}. {result['document_title']}", 1)
                            
                            # Add metadata
                            doc.add_paragraph(f"Source: {result['source_document']}")
                            doc.add_paragraph(f"Agent: {result['agent_name']}")
                            doc.add_paragraph(f"Generated: {result['generation_timestamp'][:19]}")
                            doc.add_paragraph("")  # Empty line
                            
                            # Add content
                            content = result['analysis_content']
                            markdown_to_docx(content, doc)
                            
                            if i < len(results):  # Add page break except for last document
                                doc.add_page_break()
                        
                        # Save and provide download
                        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        doc.save(temp_file.name)
                        
                        with open(temp_file.name, "rb") as file:
                            st.download_button(
                                label="Download Combined DOCX",
                                data=file.read(),
                                file_name=f"Generated_Analysis_Combined_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_combined_docx_btn"
                            )
                        
                        os.unlink(temp_file.name)
                        st.success("Combined download ready!")
                        
                    except Exception as e:
                        st.error(f"Combined download error: {e}")
            
            with col2_bulk:
                if st.button("Save All to Collection", key="save_all_to_collection"):
                    try:
                        import tempfile
                        import os
                        from io import BytesIO
                        
                        saved_count = 0
                        failed_count = 0
                        
                        with st.spinner(f"Saving {len(results)} documents to collection '{output_collection}'..."):
                            progress_bar = st.progress(0)
                            
                            for i, result in enumerate(results):
                                try:
                                    # Create temporary file
                                    temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding='utf-8')
                                    temp_file.write(result['analysis_content'])
                                    temp_file.close()
                                    
                                    # Read back as bytes
                                    with open(temp_file.name, 'rb') as file:
                                        file_data = file.read()
                                    
                                    # Create file object
                                    file_obj = BytesIO(file_data)
                                    file_obj.name = f"{result['document_title']}.txt"
                                    
                                    # Store in ChromaDB
                                    store_files_in_chromadb(
                                        [file_obj], 
                                        output_collection,
                                        model_name="basic",
                                        chunk_size=2000,
                                        chunk_overlap=200,
                                        store_images=False
                                    )
                                    
                                    saved_count += 1
                                    os.unlink(temp_file.name)
                                    
                                except Exception as e:
                                    failed_count += 1
                                    st.warning(f"Failed to save '{result['document_title']}': {e}")
                                
                                progress_bar.progress((i + 1) / len(results))
                        
                        if saved_count > 0:
                            st.success(f"Successfully saved {saved_count} documents to collection '{output_collection}'!")
                        if failed_count > 0:
                            st.warning(f"{failed_count} documents failed to save.")
                        
                    except Exception as e:
                        st.error(f"Bulk save error: {e}")
            
            with col3_bulk:
                if st.button("Clear Generated Results", key="clear_results"):
                    if st.button("Confirm Clear", key="confirm_clear"):
                        del st.session_state.generated_results
                        if 'output_collection' in st.session_state:
                            del st.session_state.output_collection
                        st.success("Generated results cleared!")
                        st.rerun()
                    else:
                        st.warning("Click 'Confirm Clear' to permanently remove results from session")
        
        else:
            # No generated results found
            st.info("No generated documents found in the current session.")
            
            # Check for existing collections with generated content
            st.subheader("Browse Existing Generated Documents")
            
            current_collections = st.session_state.collections or []
            
            if current_collections:
                # Filter for likely generated document collections
                generated_collections = [col for col in current_collections 
                                    if any(keyword in col.lower() 
                                            for keyword in ['generated', 'analysis', 'output', 'result', 'rule'])]
                
                if generated_collections:
                    st.write("**Collections that might contain generated documents:**")
                    
                    browse_collection = st.selectbox(
                        "Select collection to browse:",
                        generated_collections,
                        key="browse_generated_collection"
                    )
                    
                    if st.button("Load Generated Documents", key="load_generated_docs"):
                        try:
                            with st.spinner("Loading documents from collection..."):
                                documents = get_all_documents_in_collection(browse_collection)
                                
                                if documents:
                                    st.success(f"Found {len(documents)} documents in '{browse_collection}'")
                                    
                                    # Display documents
                                    doc_overview = []
                                    for doc in documents:
                                        doc_overview.append({
                                            "Document Name": doc["document_name"],
                                            "File Type": doc["file_type"].upper(),
                                            "Chunks": doc["total_chunks"],
                                            "Uploaded": doc["processing_timestamp"][:10] if doc["processing_timestamp"] else "Unknown",
                                            "Document ID": doc["document_id"][:12] + "..."
                                        })
                                    
                                    st.dataframe(doc_overview, use_container_width=True)
                                    
                                    # Document selector for actions
                                    doc_choices = {f"{doc['document_name']} ({doc['document_id'][:8]}...)": doc['document_id'] for doc in documents}
                                    selected_existing = st.selectbox(
                                        "Select document to view/download:",
                                        ["--Select Document--"] + list(doc_choices.keys()),
                                        key="selected_existing_doc"
                                    )
                                    
                                    if selected_existing != "--Select Document--":
                                        doc_id = doc_choices[selected_existing]
                                        
                                        col1_existing, col2_existing = st.columns(2)
                                        
                                        with col1_existing:
                                            if st.button("Preview Document", key="preview_existing_doc"):
                                                try:
                                                    with st.spinner("Loading document..."):
                                                        result = reconstruct_document_with_timeout(doc_id, browse_collection, timeout=120)
                                                        
                                                        st.subheader(f"📄 {result['document_name']}")
                                                        
                                                        content = result['reconstructed_content']
                                                        st.text_area(
                                                            f"Content ({len(content):,} characters):",
                                                            content,
                                                            height=400,
                                                            disabled=True,
                                                            key="existing_content_preview"
                                                        )
                                                        
                                                except Exception as e:
                                                    st.error(f"Preview error: {e}")
                                        
                                        with col2_existing:
                                            if st.button("Download as DOCX", key="download_existing_docx"):
                                                try:
                                                    with st.spinner("Preparing download..."):
                                                        result = reconstruct_document_with_timeout(doc_id, browse_collection, timeout=120)
                                                        
                                                        # Create DOCX
                                                        docx_path = export_to_docx(result)
                                                        
                                                        with open(docx_path, "rb") as file:
                                                            st.download_button(
                                                                label="Download DOCX",
                                                                data=file.read(),
                                                                file_name=f"{result['document_name']}.docx",
                                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                                key="download_existing_docx_btn"
                                                            )
                                                        
                                                        os.unlink(docx_path)
                                                        st.success("Download ready!")
                                                        
                                                except Exception as e:
                                                    st.error(f"Download error: {e}")
                                else:
                                    st.info(f"No documents found in collection '{browse_collection}'")
                                    
                        except Exception as e:
                            st.error(f"Error loading collection: {e}")
                
                else:
                    st.info("No collections found that appear to contain generated documents.")
            
            else:
                st.warning("No collections available.")