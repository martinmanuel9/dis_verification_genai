"""
Test Card Generation Service
Converts test procedures into executable test cards with pass/fail tracking.
Enhanced with notebook features: explicit prompts and parallel processing.
"""

from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import logging
import json

logger = logging.getLogger(__name__)

# Parallel processing configuration (from notebook)
MAX_WORKERS = 8  # Maximum concurrent test card generations

@dataclass
class TestCard:
    """Represents an executable test card"""
    test_id: str
    test_title: str
    procedures: List[str]  # Step-by-step procedures
    dependencies: List[str]
    expected_results: str
    acceptance_criteria: str
    section_title: str
    # Execution tracking (populated by users)
    executed: bool = False
    passed: bool = False
    failed: bool = False
    notes: str = ""

    def to_dict(self):
        """Convert to dictionary for JSON serialization"""
        return asdict(self)


class TestCardService:
    """
    Service for generating test cards from test plan sections.
    Based on mil_test_plan_gen.ipynb's build_test_card_for_section.
    """

    def __init__(self, llm_service):
        self.llm_service = llm_service

    def generate_test_card_from_rules(
        self,
        section_title: str,
        rules_markdown: str,
        format: str = "markdown_table"
    ) -> str:
        """
        Generate test card from test rules markdown.
        Enhanced with notebook-style explicit prompts.

        Args:
            section_title: Title of the section
            rules_markdown: Markdown containing test rules
            format: Output format ("markdown_table", "json", "docx_table")

        Returns:
            Test card in requested format
        """
        try:
            logger.info(f"Generating test card for section: {section_title} (format: {format})")

            # Extract test procedures from markdown
            test_procedures = self._extract_test_procedures_from_markdown(rules_markdown)

            if not test_procedures:
                logger.warning(f"No test procedures found in section: {section_title}")
                return self._create_empty_test_card(section_title, format)

            # Generate test card using LLM with format-specific prompts
            test_card_content = self._generate_test_card_with_llm(
                section_title,
                rules_markdown,
                test_procedures,
                format  # Pass format to LLM method
            )

            # Format output
            if format == "markdown_table":
                # For markdown, LLM returns table directly
                if isinstance(test_card_content, str):
                    return test_card_content
                return self._format_as_markdown_table(test_card_content)
            elif format == "json":
                if isinstance(test_card_content, str):
                    return test_card_content  # Already JSON string
                return self._format_as_json(test_card_content)
            elif format == "docx_table":
                return self._format_as_docx_table(test_card_content)
            else:
                raise ValueError(f"Unsupported format: {format}")

        except Exception as e:
            logger.error(f"Test card generation failed for {section_title}: {e}")
            # Return empty test card on error
            return self._create_empty_test_card(section_title, format)

    def _generate_test_card_with_llm(
        self,
        section_title: str,
        rules_markdown: str,
        test_procedures: List[Dict[str, Any]],
        format: str = "json"
    ):
        """
        Use LLM to transform test rules into structured test cards.
        Enhanced with notebook-style explicit prompts for better output quality.

        Args:
            section_title: Section title
            rules_markdown: Test rules in markdown
            test_procedures: Extracted test procedures
            format: Output format (markdown_table or json)

        Returns:
            For markdown_table: str (raw markdown table)
            For json: List[TestCard] or str (JSON string)
        """
        try:
            # Choose prompt based on format
            if format == "markdown_table":
                prompt = self._create_markdown_table_prompt(section_title, rules_markdown)
            else:  # json or docx_table
                prompt = self._create_json_prompt(section_title, rules_markdown)

            logger.info(f"Using {format} prompt for test card generation")

            response = self.llm_service.query_direct(
                model_name="gpt-4",
                query=prompt
            )[0]

            # Return based on format
            if format == "markdown_table":
                # Return raw table (notebook-style)
                logger.info(f"Generated markdown table test card")
                return response  # Raw markdown table string

            else:  # JSON format
                # Parse JSON response
                test_cards_data = self._parse_json_response(response)

                if not test_cards_data:
                    # Fallback to markdown extraction
                    logger.warning("LLM JSON parsing failed, using fallback extraction")
                    return self._extract_test_cards_from_markdown(section_title, rules_markdown)

                # Convert to TestCard objects
                test_cards = []
                for tc_data in test_cards_data:
                    test_cards.append(TestCard(
                        test_id=tc_data.get("test_id", f"TC-{len(test_cards)+1:03d}"),
                        test_title=tc_data.get("test_title", "Untitled Test"),
                        procedures=tc_data.get("procedures", []),
                        dependencies=tc_data.get("dependencies", []),
                        expected_results=tc_data.get("expected_results", ""),
                        acceptance_criteria=tc_data.get("acceptance_criteria", ""),
                        section_title=tc_data.get("section_title", section_title)
                    ))

                logger.info(f"Generated {len(test_cards)} test cards using LLM")
                return test_cards

        except Exception as e:
            logger.error(f"LLM test card generation failed: {e}")
            # Fallback: extract from markdown
            if format == "markdown_table":
                # Return empty table
                return self._create_empty_test_card(section_title, format)
            else:
                return self._extract_test_cards_from_markdown(section_title, rules_markdown)

    def _create_markdown_table_prompt(self, section_title: str, rules_markdown: str) -> str:
        """
        Create notebook-style explicit markdown table prompt.
        This produces cleaner, better-formatted tables.
        """
        return f"""You are a QA test documentation assistant.

From the following section rules (Markdown), generate a single Markdown pipe table named 'Test Card'
that lists one row per test. Do NOT include any text before or after the table.

Requirements:
- Columns: Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes
- 'Procedures' should be concise numbered steps separated by <br> (e.g., '1) Setup equipment<br>2) Configure parameters<br>3) Execute test').
- 'Expected Results' should specify exactly what happens when test passes (include specific values)
- 'Acceptance Criteria' should define measurable pass/fail thresholds
- 'Dependencies' should list prerequisites (leave empty if none)
- Leave 'Executed', 'Pass', and 'Fail' empty with checkbox (). Do NOT tick anything.
- Use short, content-based titles (not generic like "Test 1" or "Test 2")
- Derive tests from the 'Test Rules' section content

Output ONLY the table in GitHub-style pipe-table format.

=== SECTION NAME ===
{section_title}

=== SECTION RULES (MARKDOWN) ===
{rules_markdown}

=== END ==="""

    def _create_json_prompt(self, section_title: str, rules_markdown: str) -> str:
        """
        Create enhanced JSON prompt with explicit requirements.
        More detailed than original to match notebook quality.
        """
        return f"""You are a QA test documentation assistant.

From the following section rules (Markdown), generate test cards that list one test per row.

Requirements:
- Each test should be independently executable
- 'Procedures' should be concise numbered steps with explicit actions (e.g., '1) Power off equipment and disconnect from power source', '2) Connect calibrated multimeter to input terminals')
- 'Expected Results' should specify exactly what should happen when test passes (include specific values and ranges)
- 'Acceptance Criteria' should define measurable pass/fail thresholds with specific numbers
- Extract dependencies from the Dependencies section if present (specific tools, equipment, prerequisites)
- Use short, content-based titles that describe what is being tested (not generic like "Test 1")
- Derive all tests from the 'Test Rules' section

Section Name: {section_title}

Section Rules (Markdown):
{rules_markdown}

Output format: Return ONLY a JSON array of test card objects with this exact structure:
[
  {{
    "test_id": "TC-001",
    "test_title": "Brief descriptive title",
    "procedures": ["Step 1 description", "Step 2 description", "Step 3 description"],
    "expected_results": "Specific expected outcome with values and ranges",
    "acceptance_criteria": "Measurable pass/fail threshold with specific numbers",
    "dependencies": ["Specific tool or prerequisite"],
    "section_title": "{section_title}"
  }}
]

IMPORTANT: Return ONLY the JSON array, no additional text or formatting."""

    def _parse_json_response(self, response: str) -> Optional[List[Dict[str, Any]]]:
        """Parse JSON from LLM response, handling various formats"""
        try:
            # Try direct JSON parsing first
            return json.loads(response)
        except json.JSONDecodeError:
            pass

        # Try to extract JSON from markdown code blocks
        json_patterns = [
            r'```(?:json)?\s*(\[.*?\])\s*```',  # ```json [...] ```
            r'```(?:json)?\s*(\{.*?\})\s*```',  # ```json {...} ```
            r'\[.*?\]',  # Raw array
        ]

        for pattern in json_patterns:
            match = re.search(pattern, response, re.DOTALL)
            if match:
                try:
                    json_str = match.group(1) if match.lastindex else match.group(0)
                    data = json.loads(json_str)
                    if isinstance(data, list):
                        return data
                    elif isinstance(data, dict):
                        return [data]
                except json.JSONDecodeError:
                    continue

        return None

    def _format_as_markdown_table(self, test_cards: List[TestCard]) -> str:
        """
        Format test cards as GitHub-flavored markdown table.
        Enhanced with notebook-style format (includes Dependencies column).
        """
        if not test_cards:
            return "| Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes |\n|---------|------------|------------|------------------|---------------------|--------------|----------|------|------|-------|\n"

        table = "| Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes |\n"
        table += "|---------|------------|------------|------------------|---------------------|--------------|----------|------|------|-------|\n"

        for tc in test_cards:
            procedures_str = "<br>".join([f"{i+1}) {p}" for i, p in enumerate(tc.procedures)])
            # Escape pipe characters in content
            procedures_str = procedures_str.replace('|', '\\|')
            expected_results = tc.expected_results.replace('|', '\\|')
            acceptance_criteria = tc.acceptance_criteria.replace('|', '\\|')
            test_title = tc.test_title.replace('|', '\\|')
            dependencies_str = ", ".join(tc.dependencies).replace('|', '\\|') if tc.dependencies else ""

            table += f"| {tc.test_id} | {test_title} | {procedures_str} | {expected_results} | {acceptance_criteria} | {dependencies_str} |  |  |  | |\n"

        return table

    def _format_as_json(self, test_cards: List[TestCard]) -> str:
        """Format test cards as JSON"""
        return json.dumps([tc.to_dict() for tc in test_cards], indent=2)

    def _format_as_docx_table(self, test_cards: List[TestCard]) -> Any:
        """
        Format test cards as python-docx Table object.
        Returns table object that can be added to a Document.
        """
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Test ID', 'Test Title', 'Procedures', 'Expected Results',
                   'Acceptance Criteria', 'Executed', 'Pass', 'Fail', 'Notes']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Data rows
        for tc in test_cards:
            row_cells = table.add_row().cells
            row_cells[0].text = tc.test_id
            row_cells[1].text = tc.test_title
            row_cells[2].text = '\n'.join([f"{i+1}) {p}" for i, p in enumerate(tc.procedures)])
            row_cells[3].text = tc.expected_results
            row_cells[4].text = tc.acceptance_criteria
            row_cells[5].text = ''
            row_cells[6].text = ''
            row_cells[7].text = ''
            row_cells[8].text = ''

        return table

    def _extract_test_procedures_from_markdown(self, markdown: str) -> List[Dict[str, Any]]:
        """Extract test procedures from markdown (fallback method)"""
        procedures = []
        in_test_rules = False

        for line in markdown.split('\n'):
            line_stripped = line.strip()

            if line_stripped.startswith('**Test Rules:**'):
                in_test_rules = True
                continue
            elif line_stripped.startswith('**') and in_test_rules:
                # Hit another section
                break
            elif in_test_rules and re.match(r'^\d+\.', line_stripped):
                procedures.append({
                    "id": f"test_{len(procedures)+1}",
                    "description": line_stripped,
                    "type": "functional"
                })

        return procedures

    def _extract_test_cards_from_markdown(
        self,
        section_title: str,
        markdown: str
    ) -> List[TestCard]:
        """Fallback: Extract test cards directly from markdown without LLM"""
        test_cards = []
        procedures = self._extract_test_procedures_from_markdown(markdown)

        for i, proc in enumerate(procedures, 1):
            # Try to extract acceptance criteria from the description
            description = proc["description"]

            test_cards.append(TestCard(
                test_id=f"TC-{i:03d}",
                test_title=f"Test Procedure {i}",
                procedures=[description],
                dependencies=[],
                expected_results="Test completes without errors",
                acceptance_criteria="All steps pass successfully",
                section_title=section_title
            ))

        if not test_cards:
            logger.warning(f"No test cards extracted from markdown for: {section_title}")

        return test_cards

    def _create_empty_test_card(self, section_title: str, format: str) -> str:
        """Create empty test card when no procedures found"""
        if format == "markdown_table":
            return f"### Test Card: {section_title}\n\n*No testable procedures found in this section.*\n\n| Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Executed | Pass | Fail | Notes |\n|---------|------------|------------|------------------|---------------------|----------|------|------|-------|\n"
        elif format == "json":
            return "[]"
        else:
            return ""

    def generate_test_cards_for_pipeline(
        self,
        pipeline_id: str,
        redis_client,
        format: str = "markdown_table",
        max_workers: int = MAX_WORKERS
    ) -> Dict[str, str]:
        """
        Generate test cards for all sections in a pipeline with parallel processing.
        Enhanced with notebook-style parallel generation for 4-8x speedup.

        Args:
            pipeline_id: Redis pipeline ID
            redis_client: Redis client instance
            format: Output format
            max_workers: Maximum concurrent workers (default: 8)

        Returns:
            Dictionary mapping section titles to test card content
        """
        test_cards = {}

        try:
            # Get all section critic results from Redis
            pattern = f"pipeline:{pipeline_id}:critic:*"
            critic_keys = redis_client.keys(pattern)

            logger.info(f"Generating test cards for {len(critic_keys)} sections in pipeline {pipeline_id} (parallel: {max_workers} workers)")

            # Prepare tasks
            tasks = []
            for key in critic_keys:
                try:
                    critic_data = redis_client.hgetall(key)
                    section_title = critic_data.get("section_title", "")
                    synthesized_rules = critic_data.get("synthesized_rules", "")

                    if section_title and synthesized_rules:
                        tasks.append((key, section_title, synthesized_rules))
                except Exception as e:
                    logger.error(f"Failed to read data for key {key}: {e}")
                    continue

            if not tasks:
                logger.warning("No valid tasks found for test card generation")
                return {}

            # Process in parallel
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all tasks
                future_to_section = {
                    executor.submit(
                        self.generate_test_card_from_rules,
                        section_title,
                        synthesized_rules,
                        format
                    ): (key, section_title)
                    for key, section_title, synthesized_rules in tasks
                }

                # Collect results as they complete
                completed = 0
                total = len(future_to_section)
                for future in as_completed(future_to_section):
                    key, section_title = future_to_section[future]
                    completed += 1
                    try:
                        test_card = future.result()
                        test_cards[section_title] = test_card
                        logger.info(f" [{completed}/{total}] Generated test card for: {section_title}")
                    except Exception as e:
                        logger.error(f" [{completed}/{total}] Failed to generate test card for {section_title}: {e}")
                        continue

            logger.info(f"Successfully generated {len(test_cards)} test cards in parallel")
            return test_cards

        except Exception as e:
            logger.error(f"Pipeline test card generation failed: {e}")
            return {}
