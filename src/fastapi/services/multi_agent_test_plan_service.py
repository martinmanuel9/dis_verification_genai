# services/multi_agent_test_plan_service.py
"""
Multi-Agent Test Plan Generation Service - Based on mil_test_plan_gen.ipynb

Architecture:
1. Multiple Actor Agents per section using GPT-4 (can scale by adding more agents)
2. Critic Agent synthesizes actor outputs per section using GPT-4 
3. Final Critic Agent consolidates all sections using GPT-4
4. Redis Pipeline for intermediate storage and scaling
5. ChromaDB integration for section retrieval
"""

import json
import logging
import os
import re
import uuid
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import redis
import requests
from docx import Document
import base64
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import asyncio

from services.llm_service import LLMService

logger = logging.getLogger(__name__)

@dataclass
class ActorResult:
    """Result from a single actor agent"""
    agent_id: str
    model_name: str
    section_title: str
    rules_extracted: str
    processing_time: float

@dataclass
class CriticResult:
    """Result from critic agent synthesis"""
    section_title: str
    synthesized_rules: str
    dependencies: List[str]
    conflicts: List[str]
    test_procedures: List[Dict[str, Any]]
    actor_count: int

@dataclass
class FinalTestPlan:
    """Final consolidated test plan"""
    title: str
    pipeline_id: str
    total_sections: int
    total_requirements: int
    total_test_procedures: int
    consolidated_markdown: str
    processing_status: str
    sections: List[CriticResult]

class MultiAgentTestPlanService:
    def __init__(self, llm_service: LLMService, chroma_url: str, fastapi_url: str = None):
        self.llm_service = llm_service
        self.chroma_url = chroma_url.rstrip("/")
        # Use FastAPI URL for vectordb endpoints if provided, otherwise fall back to chroma_url
        self.fastapi_url = (fastapi_url or chroma_url).rstrip("/")
        
        # Redis setup for pipeline
        redis_host = os.getenv("REDIS_HOST", "redis")
        redis_port = int(os.getenv("REDIS_PORT", 6379))
        self.redis_client = redis.Redis(host=redis_host, port=redis_port, decode_responses=True)
        
        # Agent configuration using GPT-4 (scalable via env)
        # Configure actors by comma-separated ACTOR_MODELS (e.g., "gpt-4,gpt-4o-mini")
        # or by ACTOR_AGENT_COUNT (repeats "gpt-4" N times). Defaults to 3 GPT-4 actors.
        actor_models_env = os.getenv("ACTOR_MODELS")
        if actor_models_env:
            self.actor_models = [m.strip() for m in actor_models_env.split(",") if m.strip()]
        else:
            try:
                count = int(os.getenv("ACTOR_AGENT_COUNT", "3"))
            except ValueError:
                count = 3
            base_model = os.getenv("ACTOR_BASE_MODEL", "gpt-4")
            self.actor_models = [base_model for _ in range(max(1, count))]

        self.critic_model = os.getenv("CRITIC_MODEL", "gpt-4")
        self.final_critic_model = os.getenv("FINAL_CRITIC_MODEL", self.critic_model)

        # Keep originals for potential fallback logging
        self._original_actor_models = list(self.actor_models)
        self._original_critic_model = self.critic_model

        # Pipeline retention (seconds). Keep progress so UI can re-open later.
        try:
            self.pipeline_ttl_seconds = int(os.getenv("PIPELINE_TTL_SECONDS", str(60 * 60 * 24 * 7)))  # 7 days
        except ValueError:
            self.pipeline_ttl_seconds = 60 * 60 * 24 * 7
        
        logger.info(f"MultiAgentTestPlanService initialized with {len(self.actor_models)} GPT-4 actor agents")
        self._test_redis_connection()
    
    def _test_redis_connection(self):
        """Test Redis connection and setup pipeline keys"""
        try:
            self.redis_client.ping()
            logger.info("Redis connection successful")
        except Exception as e:
            logger.error(f"Redis connection failed: {e}")
    
    def generate_multi_agent_test_plan(self, 
                                     source_collections: List[str], 
                                     source_doc_ids: List[str],
                                     doc_title: str = "Test Plan") -> FinalTestPlan:
        """
        Main entry point for multi-agent test plan generation
        """
        logger.info("=== STARTING MULTI-AGENT TEST PLAN GENERATION ===")
        start_time = time.time()
        
        # Generate unique pipeline ID for this run
        pipeline_id = f"pipeline_{uuid.uuid4().hex[:12]}"
        
        try:
            # 0. Validate model availability and fallback to llama if needed
            self._maybe_fallback_to_llama(pipeline_id)

            # 1. Extract document sections from ChromaDB
            sections = self._extract_document_sections(source_collections, source_doc_ids)
            
            if not sections:
                logger.error("No sections extracted from ChromaDB")
                return self._create_fallback_test_plan(doc_title, pipeline_id)
            
            logger.info(f"Processing {len(sections)} sections with multi-agent pipeline")
            
            # 2. Initialize Redis pipeline for this run
            self._initialize_pipeline(pipeline_id, sections, doc_title)
            
            # 3. Mark pipeline as processing
            try:
                self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping={
                    "status": "PROCESSING"
                })
                self.redis_client.zadd("pipeline:processing", {pipeline_id: time.time()})
            except Exception as e:
                logger.warning(f"Failed to mark pipeline processing: {e}")

            # 4. Deploy actor agents for each section (parallel processing)
            section_results = self._deploy_section_agents(pipeline_id, sections)
            
            # 5. Deploy final critic agent to consolidate everything
            # If aborted, do not run final critic; return partial/aborted plan
            if self._is_aborted(pipeline_id):
                logger.warning(f"Pipeline {pipeline_id} aborted; skipping final critic")
                self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping={
                    "status": "ABORTED",
                    "completed_at": datetime.now().isoformat(),
                })
                self.redis_client.zrem("pipeline:processing", pipeline_id)
                aborted_markdown = f"# {doc_title}\n\nProcess aborted. {len(section_results)} sections completed before abort."
                final_plan = FinalTestPlan(
                    title=doc_title,
                    pipeline_id=pipeline_id,
                    total_sections=len(section_results),
                    total_requirements=sum(len(r.test_procedures) for r in section_results),
                    total_test_procedures=sum(len(r.test_procedures) for r in section_results),
                    consolidated_markdown=aborted_markdown,
                    processing_status="ABORTED",
                    sections=section_results,
                )
                # If purge_on_abort flag set, purge all keys except abort flag
                try:
                    meta = self.redis_client.hgetall(f"pipeline:{pipeline_id}:meta") or {}
                    if meta.get("purge_on_abort") == "1":
                        self._purge_pipeline_keys(pipeline_id)
                except Exception as e:
                    logger.warning(f"Purge on abort failed: {e}")
            else:
                final_plan = self._deploy_final_critic_agent(pipeline_id, section_results, doc_title)
            
            # 6. Mark pipeline for retention (do not hard-delete so UI can view progress)
            self._cleanup_pipeline(pipeline_id)
            # Remove from processing set
            try:
                self.redis_client.zrem("pipeline:processing", pipeline_id)
            except Exception:
                pass
            
            elapsed_time = time.time() - start_time
            logger.info(f"Multi-agent test plan generation completed in {elapsed_time:.2f}s")
            
            return final_plan
            
        except Exception as e:
            logger.error(f"Multi-agent test plan generation failed: {e}")
            self._cleanup_pipeline(pipeline_id)
            try:
                self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping={
                    "status": "FAILED",
                    "error": str(e)
                })
                self.redis_client.zrem("pipeline:processing", pipeline_id)
            except Exception:
                pass
            return self._create_fallback_test_plan(doc_title, pipeline_id)
    
    def _extract_document_sections(self, source_collections: List[str], source_doc_ids: List[str]) -> Dict[str, str]:
        """Extract sections from ChromaDB with robust reconstruction fallback.

        Strategy:
        - If explicit document IDs are provided, reconstruct full document(s) and split into natural sections.
        - Otherwise, group by metadata-based 'section_title' or page and combine chunks.
        """
        sections: Dict[str, str] = {}

        # 1) Preferred path: reconstruct by provided document IDs
        if source_doc_ids:
            for collection_name in source_collections:
                for doc_id in source_doc_ids:
                    try:
                        resp = requests.get(
                            f"{self.fastapi_url}/api/vectordb/documents/reconstruct/{doc_id}",
                            params={"collection_name": collection_name},
                            timeout=180,
                        )
                        if not resp.ok:
                            logger.warning(f"Reconstruct failed for doc_id={doc_id} in {collection_name}: {resp.status_code}")
                            continue
                        data = resp.json()
                        content = data.get("reconstructed_content") or ""
                        doc_name = data.get("document_name") or str(doc_id)
                        if content and len(content.strip()) > 50:
                            self._create_document_sections(doc_name, content, sections)
                            logger.info(f"Reconstructed {doc_name}: sections now {len(sections)}")
                    except Exception as e:
                        logger.error(f"Reconstruct error for {doc_id} in {collection_name}: {e}")

            if sections:
                logger.info(f"Extracted {len(sections)} sections via reconstruct path")
                return sections

        # 2) Fallback path: metadata grouping from the collection
        for collection_name in source_collections:
            try:
                response = requests.get(
                    f"{self.fastapi_url}/api/vectordb/documents",
                    params={"collection_name": collection_name},
                    timeout=60
                )

                if not response.ok:
                    logger.error(f"Failed to fetch from collection {collection_name}")
                    continue

                data = response.json()
                docs = data.get("documents", [])
                metas = data.get("metadatas", [])
                ids = data.get("ids", [])

                logger.info(f"Collection {collection_name} has {len(docs)} documents")

                # Group by document and section with improved metadata handling
                document_sections: Dict[str, List[str]] = {}
                for doc_id, doc, meta in zip(ids, docs, metas):
                    doc_name = (
                        (meta.get("document_name") or meta.get("filename") or meta.get("source"))
                        or doc_id
                    )

                    section_title = (
                        meta.get("section_title") or meta.get("heading") or meta.get("title")
                        or f"Section {meta.get('page_number', meta.get('page', 'Unknown'))}"
                    )

                    # Filter by explicit IDs if specified
                    if source_doc_ids:
                        if doc_id not in source_doc_ids and not any(str(x) in str(doc_name) for x in source_doc_ids):
                            continue

                    key = f"{doc_name} - {section_title}"
                    document_sections.setdefault(key, [])
                    document_sections[key].append(doc)

                # Combine chunks for each section
                for section_key, chunks in document_sections.items():
                    combined_content = "\n\n".join(chunk for chunk in chunks if chunk and chunk.strip())
                    if len(combined_content.strip()) > 100:  # Only substantial content
                        sections[section_key] = combined_content

            except Exception as e:
                logger.error(f"Error processing collection {collection_name}: {e}")

        logger.info(f"Extracted {len(sections)} sections for multi-agent processing (fallback path)")
        return sections

    def _create_document_sections(self, doc_name: str, full_document: str, sections: Dict[str, str]):
        """Create logical sections from a reconstructed full document using natural headers.

        - Prefer numbered headers, ALL-CAPS, APPENDIX/CHAPTER/SECTION markers
        - Split very large sections into sub-blocks to keep units testable
        """
        natural_sections = self._extract_natural_sections(full_document)

        # Capture distinct APPENDIX blocks if present
        appendix_pattern = re.compile(r"^APPENDIX\s+[A-Z](?:\s*[-–]\s*.*)?$", re.MULTILINE)
        appendix_matches = list(appendix_pattern.finditer(full_document))
        if appendix_matches:
            for idx, m in enumerate(appendix_matches):
                start = m.start()
                end = appendix_matches[idx + 1].start() if idx + 1 < len(appendix_matches) else len(full_document)
                title = m.group(0).strip()
                body = full_document[start:end].strip()
                if title not in natural_sections:
                    natural_sections[title] = body

        if len(natural_sections) > 1:
            for section_title, section_content in natural_sections.items():
                section_key = f"{doc_name} - {section_title}"
                content = (section_content or "").strip()
                if len(content) > 8000:
                    subsections = self._split_large_section_for_testing(content, section_title)
                    for subsection_title, subsection_content in subsections.items():
                        sections[f"{section_key} - {subsection_title}"] = subsection_content
                else:
                    sections[section_key] = content
        else:
            # Fallback to size-based split if no natural sections found
            if len(full_document) > 6000:
                parts = self._split_large_section_for_testing(full_document, "Complete Document")
                for idx, (t, c) in enumerate(parts.items(), start=1):
                    sections[f"{doc_name} - Part {idx}"] = c
            else:
                sections[f"{doc_name} - Complete Document"] = (full_document or "").strip()

    def _extract_natural_sections(self, document_text: str) -> Dict[str, str]:
        """Heuristic extraction of natural sections from text"""
        sections: Dict[str, str] = {}
        lines = document_text.split('\n')
        current_section_title = "Introduction"
        current_section_content: List[str] = []

        for line in lines:
            line_clean = line.strip()
            is_section_header = False
            header_title = ""

            if line_clean:
                # Markdown headings from reconstruction (##, ###)
                if line_clean.startswith("## ") or line_clean.startswith("### "):
                    is_section_header = True
                    # Keep original heading text
                    header_title = line_clean[3:].strip() if line_clean.startswith("## ") else line_clean[4:].strip()
                # Numbered sections like 1., 1.1, 2.3.4 followed by ALL CAPS words
                elif re.match(r'^\d+(\.\d+)*\.?\s+[A-Z][A-Z\s]+', line_clean):
                    is_section_header = True
                    header_title = line_clean
                # ALL CAPS short headers
                elif (line_clean.isupper() and len(line_clean.split()) <= 8 and len(line_clean) > 5 and not line_clean.endswith('.') and not line_clean.startswith(('THE ', 'THIS ', 'THESE '))):
                    is_section_header = True
                    header_title = line_clean
                # APPENDIX/CHAPTER/SECTION/PART
                elif line_clean.startswith(('APPENDIX', 'CHAPTER', 'SECTION', 'PART')):
                    is_section_header = True
                    header_title = line_clean
                # Short keyword headers
                elif (any(k in line_clean.upper() for k in ['REQUIREMENTS','SPECIFICATIONS','PROCEDURES','TESTING','CONFIGURATION']) and len(line_clean.split()) <= 6 and not line_clean.endswith('.') and (line_clean.startswith(tuple('0123456789')) or line_clean.isupper())):
                    is_section_header = True
                    header_title = line_clean

            if is_section_header and current_section_content:
                sections[current_section_title] = '\n'.join(current_section_content)
                current_section_title = header_title
                current_section_content = [line]
            elif is_section_header and not current_section_content:
                current_section_title = header_title
                current_section_content = [line]
            else:
                current_section_content.append(line)

        if current_section_content:
            sections[current_section_title] = '\n'.join(current_section_content)

        return sections

    def _split_large_section_for_testing(self, content: str, section_title: str) -> Dict[str, str]:
        """Split large sections into smaller units for more focused processing."""
        subsections: Dict[str, str] = {}

        # Numbered subsections first (e.g., 4.1, 4.2)
        subsection_pattern = re.compile(r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)$', re.MULTILINE)
        matches = list(subsection_pattern.finditer(content))
        if len(matches) > 1:
            for i, match in enumerate(matches):
                start = match.start()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
                subsection_num = match.group(1)
                subsection_name = match.group(2)[:50]
                subsection_content = content[start:end].strip()
                if len(subsection_content) > 500:
                    subsections[f"{subsection_num} {subsection_name}"] = subsection_content

        # Paragraph/block split fallback
        if not subsections:
            paragraphs = content.split('\n\n')
            current_block = ""
            block_num = 1
            for paragraph in paragraphs:
                if len(current_block + paragraph) > 5000:
                    if current_block.strip():
                        subsections[f"Block {block_num}"] = current_block.strip()
                        block_num += 1
                        current_block = paragraph + "\n\n"
                    else:
                        current_block += paragraph + "\n\n"
                else:
                    current_block += paragraph + "\n\n"
            if current_block.strip():
                subsections[f"Block {block_num}"] = current_block.strip()

        return subsections
    
    def _initialize_pipeline(self, pipeline_id: str, sections: Dict[str, str], doc_title: str):
        """Initialize Redis pipeline with sections and metadata"""
        pipeline_data = {
            "id": pipeline_id,
            "title": doc_title,
            "status": "INITIALIZING",
            "total_sections": len(sections),
            "sections_processed": 0,
            "created_at": datetime.now().isoformat(),
            "actor_agents": len(self.actor_models)
        }
        
        # Store pipeline metadata
        self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping=pipeline_data)
        # Track recent pipelines for quick listing in UI
        try:
            now_ts = time.time()
            self.redis_client.zadd("pipeline:recent", {pipeline_id: now_ts})
        except Exception as e:
            logger.warning(f"Failed to zadd pipeline: {e}")
        
        # Store sections for processing
        for idx, (section_title, section_content) in enumerate(sections.items()):
            section_data = {
                "title": section_title,
                "content": section_content,
                "status": "PENDING",
                "index": idx
            }
            self.redis_client.hset(f"pipeline:{pipeline_id}:section:{idx}", mapping=section_data)
        
        # Initialize result queues
        self.redis_client.delete(f"pipeline:{pipeline_id}:actor_results")
        self.redis_client.delete(f"pipeline:{pipeline_id}:critic_results")
        
        logger.info(f"Pipeline {pipeline_id} initialized with {len(sections)} sections")
    
    def _deploy_section_agents(self, pipeline_id: str, sections: Dict[str, str]) -> List[CriticResult]:
        """Deploy multiple agents per section with Redis coordination"""
        logger.info(f"Deploying agents for {len(sections)} sections")
        
        section_results = []
        
        # Process each section with multiple actor agents + critic
        with ThreadPoolExecutor(max_workers=8) as executor:
            future_to_section = {}
            
            for idx, (section_title, section_content) in enumerate(sections.items()):
                # Respect abort flag: stop submitting new work
                if self._is_aborted(pipeline_id):
                    logger.warning(f"Abort requested for pipeline {pipeline_id}; stopping new submissions at section {idx}")
                    # Mark remaining sections as aborted
                    self.redis_client.hset(f"pipeline:{pipeline_id}:section:{idx}", "status", "ABORTED")
                    break
                future = executor.submit(
                    self._process_section_with_multi_agents, 
                    pipeline_id, idx, section_title, section_content
                )
                future_to_section[future] = section_title
            
            # Collect results as they complete
            for future in as_completed(future_to_section):
                section_title = future_to_section[future]
                try:
                    critic_result = future.result(timeout=300)  # 5 minute timeout per section
                    if critic_result:
                        section_results.append(critic_result)
                        logger.info(f"Section completed: {section_title}")
                    else:
                        logger.warning(f"Section failed or aborted: {section_title}")
                except Exception as e:
                    logger.error(f"Section processing error for '{section_title}': {e}")
        
        logger.info(f"Completed processing {len(section_results)} sections")
        return section_results
    
    def _process_section_with_multi_agents(self, 
                                         pipeline_id: str, 
                                         section_idx: int,
                                         section_title: str, 
                                         section_content: str) -> Optional[CriticResult]:
        """Process a single section with multiple actor agents + critic"""
        
        # Respect abort flag early
        if self._is_aborted(pipeline_id):
            self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "ABORTED")
            return None

        # Update section status
        self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "PROCESSING")
        
        try:
            # 1. Deploy multiple actor agents in parallel
            actor_results = self._run_actor_agents(section_title, section_content)
            
            # Store actor results in Redis
            for result in actor_results:
                result_key = f"pipeline:{pipeline_id}:actor:{section_idx}:{result.agent_id}"
                result_data = {
                    "agent_id": result.agent_id,
                    "model_name": result.model_name,
                    "section_title": result.section_title,
                    "rules_extracted": result.rules_extracted,
                    "processing_time": result.processing_time
                }
                self.redis_client.hset(result_key, mapping=result_data)
            
            # 2. Deploy critic agent to synthesize actor results
            critic_result = self._run_critic_agent(section_title, section_content, actor_results)
            
            # Store critic result in Redis
            if critic_result:
                critic_key = f"pipeline:{pipeline_id}:critic:{section_idx}"
                critic_data = {
                    "section_title": critic_result.section_title,
                    "synthesized_rules": critic_result.synthesized_rules,
                    "dependencies": json.dumps(critic_result.dependencies),
                    "conflicts": json.dumps(critic_result.conflicts),
                    "test_procedures": json.dumps(critic_result.test_procedures),
                    "actor_count": critic_result.actor_count
                }
                self.redis_client.hset(critic_key, mapping=critic_data)
                
                # Update section status
                self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "COMPLETED")
                # Increment processed counter on meta
                try:
                    self.redis_client.hincrby(f"pipeline:{pipeline_id}:meta", "sections_processed", 1)
                except Exception:
                    pass
                
                return critic_result
            
        except Exception as e:
            logger.error(f"Error processing section {section_title}: {e}")
            self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "FAILED")
        
        return None

    def _is_aborted(self, pipeline_id: str) -> bool:
        try:
            return self.redis_client.get(f"pipeline:{pipeline_id}:abort") == "1"
        except Exception:
            return False
    
    def _run_actor_agents(self, section_title: str, section_content: str) -> List[ActorResult]:
        """Run multiple GPT-4 actor agents in parallel for a section"""
        actor_results = []
        
        with ThreadPoolExecutor(max_workers=len(self.actor_models)) as executor:
            futures = []
            
            for idx, model in enumerate(self.actor_models):
                agent_id = f"actor_{idx}_{uuid.uuid4().hex[:8]}"
                future = executor.submit(
                    self._run_single_actor, agent_id, model, section_title, section_content
                )
                futures.append(future)
            
            # Collect actor results
            for future in as_completed(futures):
                try:
                    result = future.result(timeout=180)  # 3 minute timeout per GPT-4 actor
                    if result:
                        actor_results.append(result)
                except Exception as e:
                    logger.error(f"Actor agent failed: {e}")
        
        logger.info(f"Completed {len(actor_results)} GPT-4 actor agents for section: {section_title}")
        return actor_results
    
    def _run_single_actor(self, agent_id: str, model: str, section_title: str, section_content: str) -> Optional[ActorResult]:
        """Run a single GPT-4 actor agent (based on notebook's extract_rules_with_llm)"""
        start_time = time.time()
        
        try:
            prompt = f"""You are a compliance and test planning expert.

Analyze the following section of a military standard and extract EVERY possible testable rule, specification, constraint, or requirement. Rules MUST be extremely detailed, explicit, and step-by-step, and should include measurable criteria, acceptable ranges, and referenced figures or tables if mentioned.

For ambiguous or implicit requirements, describe a specific test strategy.
Generate a short, content-based TITLE for this section (do not use page numbers).

ABSOLUTELY DO NOT REPEAT, DUPLICATE, OR PARAPHRASE THE SAME RULE OR LINE. Each requirement, dependency, and test step must appear ONCE ONLY.

Organize your output as follows, using markdown headings and bolded text:

## [Section Title]
**Dependencies:**
- List detailed dependencies as explicit tests, if any.

**Conflicts:**
- List detected or possible conflicts and provide recommendations or mitigation steps.

**Test Rules:**
1. (Very detailed, step-by-step numbered test rules)

Section Name: {section_title}

Section Text:
{section_content}

---
If you find truly nothing testable, reply: 'No testable rules in this section.'
"""
            
            response = self.llm_service.query_direct(
                model_name=model,
                query=prompt
            )[0]
            
            processing_time = time.time() - start_time
            
            return ActorResult(
                agent_id=agent_id,
                model_name=model,
                section_title=section_title,
                rules_extracted=response,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"GPT-4 Actor {agent_id} failed: {e}")
            return None
    
    def _run_critic_agent(self, section_title: str, section_content: str, actor_results: List[ActorResult]) -> Optional[CriticResult]:
        """Run GPT-4 critic agent to synthesize actor outputs (based on notebook's critic_review_rules)"""
        
        if not actor_results:
            logger.warning(f"No actor results to critique for section: {section_title}")
            return None
        
        try:
            # Prepare actor outputs for critic
            actor_outputs_text = ""
            for result in actor_results:
                actor_outputs_text += f"\n\nModel {result.model_name} ({result.agent_id}):\n{result.rules_extracted}\n{'='*40}"
            
            prompt = f"""You are a senior test planning reviewer (Critic AI).

Given the following section and rules extracted by several different GPT-4 models, do the following:
1. Carefully review and compare the provided rule sets.
2. Synthesize a SINGLE, detailed and explicit set of testable rules.
3. Eliminate redundancies, correct errors, and ensure all requirements are present.
4. Ensure the final test plan is step-by-step, detailed, and well organized.

NEVER simply combine all lines verbatim—synthesize, deduplicate, and streamline the content into a concise, non-repetitive format. If a rule, step, or line has the same or similar meaning as another, KEEP ONLY ONE.

Present your result in markdown format with these headings: '## [Section Title]', '**Dependencies:**', '**Conflicts:**', '**Test Rules:**'

Section Name: {section_title}

Section Text:
{section_content}

---
Actor Outputs from {len(actor_results)} GPT-4 models:
{actor_outputs_text}
"""
            
            response = self.llm_service.query_direct(
                model_name=self.critic_model,
                query=prompt
            )[0]
            
            # Apply deduplication (from notebook)
            deduplicated_response = self._deduplicate_markdown(response)
            
            # Extract structured data from critic response
            dependencies = self._extract_dependencies_from_markdown(deduplicated_response)
            conflicts = self._extract_conflicts_from_markdown(deduplicated_response)
            test_procedures = self._extract_test_procedures_from_markdown(deduplicated_response)
            
            return CriticResult(
                section_title=section_title,
                synthesized_rules=deduplicated_response,
                dependencies=dependencies,
                conflicts=conflicts,
                test_procedures=test_procedures,
                actor_count=len(actor_results)
            )
            
        except Exception as e:
            logger.error(f"GPT-4 Critic agent failed for section {section_title}: {e}")
            return None
    
    def _deploy_final_critic_agent(self, pipeline_id: str, section_results: List[CriticResult], doc_title: str) -> FinalTestPlan:
        """Deploy final GPT-4 critic agent to consolidate all sections"""
        logger.info("Deploying final GPT-4 critic agent for consolidation")
        
        try:
            # Prepare all section results for final critic
            sections_summary = []
            all_sections_content = ""
            
            for result in section_results:
                sections_summary.append({
                    "title": result.section_title,
                    "dependencies_count": len(result.dependencies),
                    "conflicts_count": len(result.conflicts),
                    "test_procedures_count": len(result.test_procedures),
                    "actor_count": result.actor_count
                })
                
                all_sections_content += f"\n\n## {result.section_title}\n"
                all_sections_content += result.synthesized_rules
                all_sections_content += "\n" + "="*60
            
            # Final critic prompt (based on notebook's final_test_plan_docx logic)
            prompt = f"""You are a final Critic AI creating a comprehensive MIL-STD test plan.

Given the following detailed section-by-section test rule reports (each synthesized from multiple GPT-4 actor agents), combine them into a single, fully ordered, professional test plan document:

1. Use a Title Page: '{doc_title}'
2. Generate a Table of Contents using ALL main section titles, in order
3. For each section, include the detailed test rules and procedures  
4. End with a 'Summary & Recommendations' section synthesizing the most critical points and overall compliance strategy

Only main content-based section titles should be in TOC, no subheadings like 'Dependencies', 'Test Rules', etc.
Preserve markdown bolds and headings for later conversion to DOCX headings and bullets.

Numbering and structure requirements:
- Produce strictly enumerated section headings: 1, 2, 3, ... (and if you introduce sub-sections, use 1.1, 1.2, 2.1, etc.)
- Start each main section heading with its number, followed by the section title
- Ensure numbering is continuous and no numbers are skipped

SECTIONS SUMMARY:
{json.dumps(sections_summary, indent=2)}

DETAILED SECTIONS:
{all_sections_content}

Create a comprehensive markdown document that consolidates all {len(section_results)} sections into a cohesive test plan.
"""
            
            response = self.llm_service.query_direct(
                model_name=self.final_critic_model,
                query=prompt
            )[0]
            
            # Apply final deduplication
            final_markdown = self._final_global_deduplicate(response)
            
            # Calculate totals
            total_requirements = sum(len(result.test_procedures) for result in section_results)
            total_test_procedures = total_requirements  # Each requirement becomes a test procedure
            
            # Store final result in Redis
            final_result_key = f"pipeline:{pipeline_id}:final_result"
            final_data = {
                "title": doc_title,
                "consolidated_markdown": final_markdown,
                "total_sections": len(section_results),
                "total_requirements": total_requirements,
                "total_test_procedures": total_test_procedures,
                "processing_status": "COMPLETED",
                "completed_at": datetime.now().isoformat()
            }
            self.redis_client.hset(final_result_key, mapping=final_data)
            # Update pipeline meta status and bump recency
            try:
                self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping={
                    "status": "COMPLETED",
                    "completed_at": final_data["completed_at"],
                })
                self.redis_client.zadd("pipeline:recent", {pipeline_id: time.time()})
            except Exception as e:
                logger.warning(f"Failed to update pipeline meta completion: {e}")
            
            return FinalTestPlan(
                title=doc_title,
                pipeline_id=pipeline_id,
                total_sections=len(section_results),
                total_requirements=total_requirements,
                total_test_procedures=total_test_procedures,
                consolidated_markdown=final_markdown,
                processing_status="COMPLETED",
                sections=section_results
            )
            
        except Exception as e:
            logger.error(f"Final GPT-4 critic agent failed: {e}")
            return FinalTestPlan(
                title=doc_title,
                pipeline_id=pipeline_id,
                total_sections=len(section_results),
                total_requirements=0,
                total_test_procedures=0,
                consolidated_markdown=f"# {doc_title}\n\nFinal critic agent failed: {str(e)}",
                processing_status="FAILED",
                sections=section_results
            )
    
    def _deduplicate_markdown(self, text: str) -> str:
        """Deduplicate sentences within markdown sections (from notebook)"""
        output = []
        section_boundary = lambda l: l.startswith("## ") or (l.startswith("**") and l.endswith("**"))
        
        def process_block(block):
            local_seen = set()
            for sentence in re.split(r'(?<=[.!?]) +', block):
                sent = sentence.strip()
                norm = re.sub(r'\s+', ' ', sent.lower())
                if not sent or norm in local_seen:
                    continue
                output.append(sent)
                local_seen.add(norm)
        
        current_block = []
        for line in text.split('\n'):
            if section_boundary(line):
                process_block(' '.join(current_block))
                current_block = []
                output.append(line)
            elif line.strip() == "":
                process_block(' '.join(current_block))
                current_block = []
                output.append(line)
            else:
                current_block.append(line.strip())
        
        process_block(' '.join(current_block))
        return '\n'.join(output)
    
    def _final_global_deduplicate(self, text: str) -> str:
        """Remove duplicate lines globally (from notebook)"""
        seen = set()
        out = []
        
        for line in text.split('\n'):
            sentences = re.split(r'(?<=[.!?]) +', line) if len(line) > 120 else [line]
            unique_sentences = []
            
            for s in sentences:
                norm = re.sub(r'\s+', ' ', s.strip().lower())
                if norm and norm not in seen:
                    unique_sentences.append(s)
                    seen.add(norm)
                elif not s.strip():
                    unique_sentences.append(s)
            
            joined = ' '.join(unique_sentences).strip()
            if joined or not line.strip():
                out.append(joined)
        
        return '\n'.join(out)
    
    def _extract_dependencies_from_markdown(self, markdown: str) -> List[str]:
        """Extract dependencies from markdown format"""
        dependencies = []
        in_dependencies = False
        
        for line in markdown.split('\n'):
            if line.strip().startswith('**Dependencies:**'):
                in_dependencies = True
                continue
            elif line.strip().startswith('**') and in_dependencies:
                break
            elif in_dependencies and line.strip().startswith('- '):
                dependencies.append(line.strip()[2:])
        
        return dependencies
    
    def _extract_conflicts_from_markdown(self, markdown: str) -> List[str]:
        """Extract conflicts from markdown format"""
        conflicts = []
        in_conflicts = False
        
        for line in markdown.split('\n'):
            if line.strip().startswith('**Conflicts:**'):
                in_conflicts = True
                continue
            elif line.strip().startswith('**') and in_conflicts:
                break
            elif in_conflicts and line.strip().startswith('- '):
                conflicts.append(line.strip()[2:])
        
        return conflicts
    
    def _extract_test_procedures_from_markdown(self, markdown: str) -> List[Dict[str, Any]]:
        """Extract test procedures from markdown format"""
        procedures = []
        in_test_rules = False
        
        for line in markdown.split('\n'):
            if line.strip().startswith('**Test Rules:**'):
                in_test_rules = True
                continue
            elif line.strip().startswith('**') and in_test_rules:
                break
            elif in_test_rules and re.match(r'^\d+\.', line.strip()):
                procedures.append({
                    "id": f"test_{len(procedures)+1}",
                    "description": line.strip(),
                    "type": "functional"
                })
        
        return procedures
    
    def _cleanup_pipeline(self, pipeline_id: str):
        """Mark pipeline keys with an expiration instead of hard deletion so UI can inspect later."""
        try:
            pattern = f"pipeline:{pipeline_id}:*"
            keys = self.redis_client.keys(pattern)
            for k in keys:
                try:
                    self.redis_client.expire(k, self.pipeline_ttl_seconds)
                except Exception:
                    pass
            # Also keep the meta record updated and expiring
            self.redis_client.expire(f"pipeline:{pipeline_id}:meta", self.pipeline_ttl_seconds)
            logger.info(f"Retained {len(keys)} Redis keys for pipeline {pipeline_id} with TTL={self.pipeline_ttl_seconds}s")
        except Exception as e:
            logger.error(f"Error retaining pipeline {pipeline_id}: {e}")

    def _purge_pipeline_keys(self, pipeline_id: str):
        try:
            pattern = f"pipeline:{pipeline_id}:*"
            keys = self.redis_client.keys(pattern) or []
            keep_key = f"pipeline:{pipeline_id}:abort"
            to_delete = [k for k in keys if k != keep_key]
            if to_delete:
                self.redis_client.delete(*to_delete)
            # Remove from index sets
            try:
                self.redis_client.zrem("pipeline:recent", pipeline_id)
                self.redis_client.zrem("pipeline:processing", pipeline_id)
            except Exception:
                pass
            logger.info(f"Purged pipeline {pipeline_id} keys: {len(to_delete)} deleted")
        except Exception as e:
            logger.error(f"Error purging pipeline {pipeline_id}: {e}")

    # ===========================
    # Model availability + fallback
    # ===========================
    def _maybe_fallback_to_llama(self, pipeline_id: str):
        """If OpenAI models are configured but unavailable or quota-exceeded, switch to llama models.

        Strategy:
        - If any actor/critic model starts with 'gpt' but OPEN_AI_API_KEY is missing -> fallback to llama
        - Else, attempt a very small direct query with critic model; on exception -> fallback to llama
        - Record fallback in Redis pipeline meta
        """
        try:
            needs_openai = any(str(m).lower().startswith("gpt") for m in (self.actor_models + [self.critic_model, self.final_critic_model]))
            if not needs_openai:
                return

            openai_key = os.getenv("OPENAI_API_KEY")
            if not openai_key:
                self._apply_llama_fallback(pipeline_id, reason="missing_openai_key")
                return

            # Quick probe with critic model
            try:
                # Keep it very short; do not log history
                _ = self.llm_service.query_direct(self.critic_model, "Return OK.", session_id=None, log_history=False)
                # If success, keep current models
                return
            except Exception as e:
                # Any failure here triggers fallback
                self._apply_llama_fallback(pipeline_id, reason=f"probe_failed: {str(e)[:120]}")
        except Exception as e:
            # Never fail generation due to fallback logic
            logger.warning(f"Model fallback check error: {e}")

    def _apply_llama_fallback(self, pipeline_id: str, reason: str = "unavailable"):
        try:
            # Use Llama 3.1 8B via Ollama for local CPU-based fallback
            fallback_model = "llama3.1:8b"
            logger.warning(f"Falling back to {fallback_model} models due to: {reason}")
            self.actor_models = [fallback_model for _ in self.actor_models]
            self.critic_model = fallback_model
            self.final_critic_model = fallback_model
            # Record in Redis
            self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping={
                "model_fallback": fallback_model,
                "fallback_reason": reason,
                "original_actor_models": json.dumps(self._original_actor_models),
                "original_critic_model": self._original_critic_model,
                "actor_agents": len(self.actor_models),
            })
        except Exception as e:
            logger.warning(f"Failed to record model fallback: {e}")
    
    def _create_fallback_test_plan(self, doc_title: str, pipeline_id: str) -> FinalTestPlan:
        """Create fallback test plan"""
        fallback_markdown = f"""# {doc_title}

## Fallback Mode Notice
This test plan was generated in fallback mode due to section extraction issues.

## Multi-Agent Architecture
- 3x GPT-4 Actor Agents per section
- 1x GPT-4 Critic Agent per section  
- 1x GPT-4 Final Critic Agent for consolidation
- Redis pipeline for scaling

## Next Steps
1. Fix ChromaDB section extraction
2. Verify collection names and document IDs
3. Re-run multi-agent pipeline with GPT-4 agents
"""
        
        return FinalTestPlan(
            title=doc_title,
            pipeline_id=pipeline_id,
            total_sections=0,
            total_requirements=0,
            total_test_procedures=0,
            consolidated_markdown=fallback_markdown,
            processing_status="FALLBACK",
            sections=[]
        )
    
    def export_to_word(self, test_plan: FinalTestPlan) -> str:
        """Export test plan to Word document"""
        doc = Document()
        doc.add_heading(test_plan.title, level=0)
        
        # Add multi-agent architecture info
        doc.add_heading('Multi-Agent Architecture', level=2)
        doc.add_paragraph("This test plan was generated using:")
        # Use Word list formatting only; do not include literal bullet characters
        doc.add_paragraph("3x GPT-4 Actor Agents per section", style='List Bullet')
        doc.add_paragraph("1x GPT-4 Critic Agent per section for synthesis", style='List Bullet')  
        doc.add_paragraph("1x GPT-4 Final Critic Agent for consolidation", style='List Bullet')
        doc.add_paragraph("Redis pipeline for scalable processing", style='List Bullet')
        
        # Convert markdown to Word
        self._convert_markdown_to_docx(test_plan.consolidated_markdown, doc)
        
        # Add statistics
        doc.add_heading('Test Plan Statistics', level=2)
        doc.add_paragraph(f"Total sections processed: {test_plan.total_sections}")
        doc.add_paragraph(f"Total requirements: {test_plan.total_requirements}")
        doc.add_paragraph(f"Total test procedures: {test_plan.total_test_procedures}")
        doc.add_paragraph(f"Processing status: {test_plan.processing_status}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Serialize and encode
        buf = io.BytesIO()
        doc.save(buf)
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    
    def _convert_markdown_to_docx(self, markdown_content: str, doc: Document):
        """Convert markdown to Word document (based on notebook's markdown_to_docx)"""
        lines = markdown_content.split('\n')
        
        for line in lines:
            l = line.strip()
            if not l:
                continue
                
            if l.startswith('# '):
                continue  # Skip main title as already added
            elif l.startswith('## '):  # Section/subsection heading
                doc.add_heading(l.replace("##", "").strip(), level=1)
            elif l.startswith("**") and l.endswith("**"):
                doc.add_heading(l.replace("*", "").strip(), level=2)
            elif l.startswith(("-", "*", "•")):
                doc.add_paragraph(l.lstrip("-*• ").strip(), style='List Bullet')
            elif l[:2].isdigit() and l[2] in ('.', ')'):
                doc.add_paragraph(l, style='List Number')
            elif "**" in l:  # Remove inline bold
                parts = l.split("**")
                p = doc.add_paragraph()
                toggle = False
                for part in parts:
                    run = p.add_run(part)
                    if toggle:
                        run.bold = True
                    toggle = not toggle
            else:
                doc.add_paragraph(l)
    
    def save_to_chromadb(self, test_plan: FinalTestPlan, session_id: str, pipeline_id: Optional[str] = None) -> Dict[str, Any]:
        """Save generated test plan to ChromaDB (single collection) with idempotency per pipeline.

        If a generated_document_id is already recorded in pipeline meta, re-use it and do not create a duplicate.
        """
        try:
            # First, ensure the target collection exists
            target_collection = os.getenv("GENERATED_TESTPLAN_COLLECTION", "generated_test_plan")
            self._ensure_collection_exists(target_collection)

            # Determine pipeline id
            pid = pipeline_id or getattr(test_plan, 'pipeline_id', None)

            # If a document was already saved for this pipeline, return that
            if pid:
                try:
                    meta = self.redis_client.hgetall(f"pipeline:{pid}:meta") or {}
                    existing_id = meta.get("generated_document_id")
                    if existing_id:
                        return {
                            "document_id": existing_id,
                            "collection_name": meta.get("collection", target_collection),
                            "saved": True,
                            "generated_at": meta.get("completed_at") or datetime.now().isoformat(),
                            "architecture": "multi_agent_gpt4",
                            "reused": True
                        }
                except Exception:
                    pass

                # Acquire a short-lived lock to avoid concurrent double-save
                try:
                    lock_key = f"pipeline:{pid}:saving_lock"
                    if not self.redis_client.set(lock_key, "1", nx=True, ex=60):
                        # Someone else is saving; wait briefly then return recorded id if present
                        time.sleep(1)
                        meta = self.redis_client.hgetall(f"pipeline:{pid}:meta") or {}
                        existing_id = meta.get("generated_document_id")
                        if existing_id:
                            return {
                                "document_id": existing_id,
                                "collection_name": meta.get("collection", target_collection),
                                "saved": True,
                                "generated_at": meta.get("completed_at") or datetime.now().isoformat(),
                                "architecture": "multi_agent_gpt4",
                                "reused": True
                            }
                except Exception:
                    pass

            # Create document content
            doc_content = f"Title: {test_plan.title}\n\n"
            doc_content += f"Architecture: Multi-Agent GPT-4 Pipeline\n"
            doc_content += f"Sections: {test_plan.total_sections}\n"
            doc_content += f"Requirements: {test_plan.total_requirements}\n" 
            doc_content += f"Test Procedures: {test_plan.total_test_procedures}\n\n"
            doc_content += test_plan.consolidated_markdown
            
            # Prepare metadata
            metadata = {
                "title": test_plan.title,
                "type": "generated_test_plan",
                "architecture": "multi_agent_gpt4",
                "session_id": session_id,
                "generated_at": datetime.now().isoformat(),
                "total_sections": test_plan.total_sections,
                "total_requirements": test_plan.total_requirements,
                "total_test_procedures": test_plan.total_test_procedures,
                "processing_status": test_plan.processing_status,
                "agent_types": "3x_gpt4_actors_1x_gpt4_critic_1x_final_critic",
                "word_count": len(doc_content.split()),
                "char_count": len(doc_content)
            }
            
            # Generate unique document ID
            doc_id = f"testplan_multiagent_{(pid or session_id)}_{uuid.uuid4().hex[:8]}"
            
            # Save to ChromaDB
            payload = {
                "collection_name": target_collection, 
                "documents": [doc_content],
                "metadatas": [metadata],
                "ids": [doc_id]
            }
            
            response = requests.post(
                f"{self.fastapi_url}/api/vectordb/documents/add",
                json=payload,
                timeout=30
            )
            
            if response.ok:
                logger.info(f"Saved multi-agent test plan to ChromaDB ({target_collection}): {doc_id}")
                # Record generated doc in pipeline meta if pipeline_id provided
                try:
                    if pid:
                        self.redis_client.hset(f"pipeline:{pid}:meta", mapping={
                            "generated_document_id": doc_id,
                            "collection": target_collection,
                        })
                        self.redis_client.zadd("pipeline:recent", {pid: time.time()})
                except Exception as e:
                    logger.warning(f"Failed to record generated doc in pipeline meta: {e}")

                return {
                    "document_id": doc_id,
                    "collection_name": target_collection,
                    "saved": True,
                    "generated_at": metadata["generated_at"],
                    "architecture": "multi_agent_gpt4"
                }
            else:
                logger.error(f"Failed to save to ChromaDB: {response.status_code} - {response.text}")
                return {"saved": False, "error": response.text}

        except Exception as e:
            logger.error(f"Error saving to ChromaDB: {e}")
            return {"saved": False, "error": str(e)}
    
    def _ensure_generated_documents_collection_exists(self):
        """Deprecated no-op: use _ensure_collection_exists with GENERATED_TESTPLAN_COLLECTION instead."""
        try:
            self._ensure_collection_exists(os.getenv("GENERATED_TESTPLAN_COLLECTION", "generated_test_plan"))
        except Exception:
            pass

    def _ensure_collection_exists(self, name: str):
        """Ensure an arbitrary collection exists."""
        try:
            list_response = requests.get(f"{self.fastapi_url}/api/vectordb/collections", timeout=10)
            if list_response.ok:
                collections = list_response.json().get("collections", [])
                if name not in collections:
                    requests.post(
                        f"{self.fastapi_url}/api/vectordb/collection/create",
                        params={"collection_name": name},
                        timeout=10
                    )
        except Exception as e:
            logger.warning(f"Unable to ensure collection {name}: {e}")
