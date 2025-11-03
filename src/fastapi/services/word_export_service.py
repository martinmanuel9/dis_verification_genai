"""
Word Export Service
Provides Word document export capabilities for agents, chat history, and agent responses.
"""

import os
import base64
import uuid
from datetime import datetime, timezone
from typing import Dict, List, Any, Optional, Union
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from sqlalchemy.orm import Session
from services.database import ComplianceAgent, ChatHistory, AgentResponse, DebateSession
import logging

logger = logging.getLogger("WORD_EXPORT_SERVICE")

class WordExportService:
    """
    Service for exporting various application data to Word documents.
    Enhanced with streaming capabilities and performance optimizations.
    """
    
    def __init__(self, enable_streaming: bool = True, chunk_size: int = 50):
        self.temp_dir = os.path.join(os.getcwd(), "temp_exports")
        os.makedirs(self.temp_dir, exist_ok=True)
        self.enable_streaming = enable_streaming
        self.chunk_size = chunk_size  # Number of items to process per chunk
        self._doc_cache = {}  # Cache for frequently used document styles
    
    def export_agents_to_word(self, agents: List[Dict[str, Any]], export_format: str = "detailed") -> bytes:
        """
        Export agent configurations to a Word document.
        Enhanced with streaming for large datasets.
        
        Args:
            agents: List of agent dictionaries
            export_format: "summary" or "detailed"
        
        Returns:
            bytes: Word document content
        """
        try:
            start_time = datetime.now()
            logger.info(f"Starting Word export for {len(agents)} agents in {export_format} format")
            
            doc = Document()
            
            # Set up document styles with caching
            self._setup_document_styles_optimized(doc)
            
            # Title
            title = doc.add_heading('Agent Configurations Export', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Total Agents: {len(agents)}")
            doc.add_paragraph(f"Export Format: {export_format.title()}")
            doc.add_paragraph("")
            
            # Agents section
            doc.add_heading('Agent Details', level=1)
            
            # Process agents in chunks for better memory management
            if self.enable_streaming and len(agents) > self.chunk_size:
                return self._export_agents_streaming(doc, agents, export_format)
            else:
                return self._export_agents_traditional(doc, agents, export_format)
                
        except Exception as e:
            logger.error(f"Failed to export agents to Word: {str(e)}")
            raise e
    
    def _export_agents_streaming(self, doc: Document, agents: List[Dict[str, Any]], export_format: str) -> bytes:
        """Export agents using streaming approach for better performance"""
        logger.info(f"Using streaming export for {len(agents)} agents")
        
        chunks = [agents[i:i + self.chunk_size] for i in range(0, len(agents), self.chunk_size)]
        
        for chunk_idx, chunk in enumerate(chunks):
            logger.debug(f"Processing chunk {chunk_idx + 1}/{len(chunks)} ({len(chunk)} agents)")
            
            for i, agent in enumerate(chunk, chunk_idx * self.chunk_size + 1):
                self._add_agent_to_document(doc, agent, i, export_format)
                
                # Add page break between agents (except last one)
                if i < len(agents):
                    doc.add_page_break()
        
        # Convert to bytes
        doc_bytes = self._document_to_bytes(doc)
        logger.info(f"Streaming export completed for {len(agents)} agents")
        return doc_bytes
    
    def _export_agents_traditional(self, doc: Document, agents: List[Dict[str, Any]], export_format: str) -> bytes:
        """Traditional export approach for smaller datasets"""
        logger.debug(f"Using traditional export for {len(agents)} agents")
        
        for i, agent in enumerate(agents, 1):
            self._add_agent_to_document(doc, agent, i, export_format)
            
            # Page break between agents (except last one)
            if i < len(agents):
                doc.add_page_break()
        
        # Convert to bytes
        doc_bytes = self._document_to_bytes(doc)
        logger.info(f"Traditional export completed for {len(agents)} agents")
        return doc_bytes
    
    def _add_agent_to_document(self, doc: Document, agent: Dict[str, Any], agent_num: int, export_format: str):
        """Add a single agent to the document (optimized)"""
        # Agent header
        agent_heading = doc.add_heading(f"{agent_num}. {agent.get('name', 'Unknown Agent')}", level=2)
        
        # Basic information table (optimized)
        self._add_agent_basic_info_optimized(doc, agent)
        
        if export_format == "detailed":
            # System prompt
            system_prompt = agent.get('system_prompt', 'No system prompt defined')
            if system_prompt and len(system_prompt.strip()) > 0:
                doc.add_heading('System Prompt', level=3)
                prompt_para = doc.add_paragraph(system_prompt[:1000] + "..." if len(system_prompt) > 1000 else system_prompt)
                prompt_para.style = 'Quote'
            
            # User prompt template
            user_prompt = agent.get('user_prompt_template', 'No user prompt template defined')
            if user_prompt and len(user_prompt.strip()) > 0:
                doc.add_heading('User Prompt Template', level=3)
                template_para = doc.add_paragraph(user_prompt[:1000] + "..." if len(user_prompt) > 1000 else user_prompt)
                template_para.style = 'Quote'
            
            # Performance metrics if available
            if agent.get('total_queries', 0) > 0:
                doc.add_heading('Performance Metrics', level=3)
                self._add_agent_performance_metrics_optimized(doc, agent)
    
    def export_chat_history_to_word(self, chat_history: List[Dict[str, Any]], session_filter: Optional[str] = None) -> bytes:
        """
        Export chat history to a Word document.
        
        Args:
            chat_history: List of chat history records
            session_filter: Optional session ID to filter by
        
        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            title = doc.add_heading('Chat History Export', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Total Conversations: {len(chat_history)}")
            if session_filter:
                doc.add_paragraph(f"Session Filter: {session_filter}")
            doc.add_paragraph("")
            
            # Group by session if multiple sessions
            sessions = {}
            for chat in chat_history:
                session_id = chat.get('session_id', 'unknown')
                if session_id not in sessions:
                    sessions[session_id] = []
                sessions[session_id].append(chat)
            
            # Export each session
            for session_id, chats in sessions.items():
                doc.add_heading(f'Session: {session_id}', level=1)
                
                # Session summary
                session_start = min(chat.get('timestamp', datetime.now()) for chat in chats)
                session_end = max(chat.get('timestamp', datetime.now()) for chat in chats)
                
                doc.add_paragraph(f"Session Duration: {session_start} to {session_end}")
                doc.add_paragraph(f"Total Interactions: {len(chats)}")
                doc.add_paragraph("")
                
                # Individual chats
                for i, chat in enumerate(chats, 1):
                    self._add_chat_interaction(doc, chat, i)
                
                # Page break between sessions
                if len(sessions) > 1 and session_id != list(sessions.keys())[-1]:
                    doc.add_page_break()
            
            doc_bytes = self._document_to_bytes(doc)
            logger.info(f"Exported {len(chat_history)} chat records to Word document")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to export chat history to Word: {str(e)}")
            raise e
    
    def export_agent_simulation_to_word(self, simulation_data: Dict[str, Any]) -> bytes:
        """
        Export agent simulation results to a Word document.
        
        Args:
            simulation_data: Dictionary containing simulation results
        
        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            title = doc.add_heading('AI Agent Simulation Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            if simulation_data.get('session_id'):
                doc.add_paragraph(f"Session ID: {simulation_data['session_id']}")
            doc.add_paragraph("")
            
            # Simulation overview
            doc.add_heading('Simulation Overview', level=1)
            
            simulation_type = simulation_data.get('type', 'Unknown')
            doc.add_paragraph(f"Simulation Type: {simulation_type}")
            
            if 'query' in simulation_data:
                doc.add_heading('Original Query/Content', level=2)
                query_para = doc.add_paragraph(simulation_data['query'])
                query_para.style = 'Quote'
            
            # Agents involved
            if 'agents' in simulation_data:
                doc.add_heading('Participating Agents', level=2)
                agents_table = doc.add_table(rows=1, cols=3)
                agents_table.style = 'Table Grid'
                
                # Header
                header_cells = agents_table.rows[0].cells
                header_cells[0].text = 'Agent Name'
                header_cells[1].text = 'Model'
                header_cells[2].text = 'Role'
                
                for agent in simulation_data['agents']:
                    row_cells = agents_table.add_row().cells
                    row_cells[0].text = agent.get('name', 'Unknown')
                    row_cells[1].text = agent.get('model_name', 'Unknown')
                    row_cells[2].text = agent.get('role', 'Analysis')
            
            # Results section
            doc.add_heading('Simulation Results', level=1)
            
            if 'agent_responses' in simulation_data:
                self._add_agent_responses(doc, simulation_data['agent_responses'])
            elif 'debate_chain' in simulation_data:
                self._add_debate_chain(doc, simulation_data['debate_chain'])
            elif 'details' in simulation_data:
                self._add_compliance_details(doc, simulation_data['details'])
            
            # Performance metrics if available
            if simulation_data.get('response_time_ms'):
                doc.add_heading('Performance Metrics', level=2)
                doc.add_paragraph(f"Total Response Time: {simulation_data['response_time_ms']/1000:.2f} seconds")
            
            doc_bytes = self._document_to_bytes(doc)
            logger.info("Exported agent simulation to Word document")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to export agent simulation to Word: {str(e)}")
            raise e
    
    def export_rag_assessment_to_word(self, assessment_data: Dict[str, Any]) -> bytes:
        """
        Export RAG assessment results to a Word document.
        
        Args:
            assessment_data: Dictionary containing assessment results
        
        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            title = doc.add_heading('RAG Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            if assessment_data.get('performance_metrics', {}).get('session_id'):
                doc.add_paragraph(f"Session ID: {assessment_data['performance_metrics']['session_id']}")
            doc.add_paragraph("")
            
            # Original query
            if 'performance_metrics' in assessment_data and 'query' in assessment_data['performance_metrics']:
                doc.add_heading('Original Query', level=1)
                query_para = doc.add_paragraph(assessment_data['performance_metrics']['query'])
                query_para.style = 'Quote'
            
            # Generated response
            if 'response' in assessment_data:
                doc.add_heading('Generated Response', level=1)
                response_para = doc.add_paragraph(assessment_data['response'])
                response_para.style = 'Quote'
            
            # Performance metrics
            if 'performance_metrics' in assessment_data:
                doc.add_heading('Performance Metrics', level=1)
                self._add_performance_metrics_table(doc, assessment_data['performance_metrics'])
            
            # Quality assessment
            if 'quality_assessment' in assessment_data and assessment_data['quality_assessment']:
                doc.add_heading('Quality Assessment', level=1)
                self._add_quality_assessment_table(doc, assessment_data['quality_assessment'])
            
            # Alignment assessment
            if 'alignment_assessment' in assessment_data and assessment_data['alignment_assessment']:
                doc.add_heading('Alignment Assessment', level=1)
                self._add_alignment_assessment_table(doc, assessment_data['alignment_assessment'])
            
            # Classification metrics
            if 'classification_metrics' in assessment_data and assessment_data['classification_metrics']:
                doc.add_heading('Classification Metrics', level=1)
                self._add_classification_metrics_table(doc, assessment_data['classification_metrics'])
            
            doc_bytes = self._document_to_bytes(doc)
            logger.info("Exported RAG assessment to Word document")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to export RAG assessment to Word: {str(e)}")
            raise e
    
    def export_reconstructed_document_to_word(self, reconstructed: Dict[str, Any]) -> bytes:
        """Export a reconstructed document (text + optional images) to a Word document."""
        try:
            doc = Document()
            self._setup_document_styles(doc)

            title_text = reconstructed.get('document_name') or 'Reconstructed Document'
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Render reconstructed content with simple markdown-like handling
            content = reconstructed.get('reconstructed_content', '') or ''
            for line in content.splitlines():
                l = (line or '').strip()
                if not l:
                    doc.add_paragraph("")
                    continue
                if l.startswith('### '):
                    doc.add_heading(l[4:].strip(), level=3)
                elif l.startswith('## '):
                    doc.add_heading(l[3:].strip(), level=2)
                elif l.startswith('# '):
                    doc.add_heading(l[2:].strip(), level=1)
                elif l.startswith(('-', '*', '•')):
                    doc.add_paragraph(l.lstrip('-*• ').strip(), style='List Bullet')
                else:
                    doc.add_paragraph(l)

            # Optional: embed images
            images = reconstructed.get('images') or []
            if images:
                doc.add_page_break()
                doc.add_heading('Images', level=1)
                CHROMA_URL = os.getenv("CHROMA_URL", "http://localhost:8000")
                import tempfile
                from docx.shared import Inches
                for idx, img in enumerate(images, 1):
                    filename = img.get('filename')
                    desc = img.get('description') or ''
                    try:
                        import requests
                        resp = requests.get(f"{CHROMA_URL}/images/{filename}", timeout=15)
                        if resp.ok:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tf:
                                tf.write(resp.content)
                                temp_path = tf.name
                            doc.add_picture(temp_path, width=Inches(5.5))
                            if desc:
                                para = doc.add_paragraph(desc)
                                para.style = 'Quote'
                            os.unlink(temp_path)
                        else:
                            doc.add_paragraph(f"[Image '{filename}' not available]")
                            if desc:
                                para = doc.add_paragraph(desc)
                                para.style = 'Quote'
                    except Exception as e:
                        doc.add_paragraph(f"[Failed to load image '{filename}': {e}]")
                        if desc:
                            para = doc.add_paragraph(desc)
                            para.style = 'Quote'

            return self._document_to_bytes(doc)
        except Exception as e:
            logger.error(f"Failed to export reconstructed document to Word: {str(e)}")
            raise e

    def export_markdown_to_word(self, title: str, markdown_content: str) -> bytes:
        """Export generic markdown-like content to a Word document."""
        try:
            doc = Document()
            self._setup_document_styles(doc)

            # Title
            if title:
                t = doc.add_heading(title, 0)
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Basic markdown handling
            for line in (markdown_content or '').splitlines():
                l = (line or '').strip()
                if not l:
                    doc.add_paragraph("")
                    continue
                if l.startswith('### '):
                    doc.add_heading(l[4:].strip(), level=3)
                elif l.startswith('## '):
                    doc.add_heading(l[3:].strip(), level=2)
                elif l.startswith('# '):
                    doc.add_heading(l[2:].strip(), level=1)
                elif l.startswith(('-', '*', '•')):
                    doc.add_paragraph(l.lstrip('-*• ').strip(), style='List Bullet')
                elif l[:2].isdigit() and len(l) > 2 and l[2] in ('.', ')'):
                    doc.add_paragraph(l, style='List Number')
                else:
                    # Handle inline bold via **text**
                    if '**' in l:
                        parts = l.split('**')
                        p = doc.add_paragraph()
                        bold = False
                        for part in parts:
                            run = p.add_run(part)
                            if bold:
                                run.bold = True
                            bold = not bold
                    else:
                        doc.add_paragraph(l)

            return self._document_to_bytes(doc)
        except Exception as e:
            logger.error(f"Failed to export markdown to Word: {str(e)}")
            raise e
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document."""
        try:
            # Create custom styles
            styles = doc.styles
            
            # Quote style
            if 'Quote' not in [style.name for style in styles]:
                quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
                quote_font = quote_style.font
                quote_font.italic = True
                quote_font.size = Pt(10)
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.right_indent = Inches(0.5)
        except:
            pass  # Ignore styling errors
    
    def _setup_document_styles_optimized(self, doc: Document):
        """Set up custom styles for the document with caching."""
        doc_id = id(doc)
        
        if doc_id in self._doc_cache:
            logger.debug("Using cached document styles")
            return
        
        try:
            # Create custom styles
            styles = doc.styles
            
            # Quote style
            if 'Quote' not in [style.name for style in styles]:
                quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
                quote_font = quote_style.font
                quote_font.italic = True
                quote_font.size = Pt(10)
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.right_indent = Inches(0.5)
                
                # Cache the setup
                self._doc_cache[doc_id] = True
                logger.debug("Document styles cached")
        except Exception as e:
            logger.warning(f"Failed to set up document styles: {e}")
    
    def _add_agent_basic_info_optimized(self, doc: Document, agent: Dict[str, Any]):
        """Add basic agent information as a table (optimized version)."""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Prepare data more efficiently
        created_at = agent.get('created_at', 'Unknown')
        if isinstance(created_at, str) and len(created_at) > 10:
            created_at = created_at[:10]  # Just date part
        
        rows_data = [
            ('Agent ID', str(agent.get('id', 'N/A'))),
            ('Model', agent.get('model_name', 'Unknown')),
            ('Created', created_at),
            ('Status', 'Active' if agent.get('is_active', True) else 'Inactive'),
            ('Temperature', str(agent.get('temperature', 'N/A'))),
            ('Max Tokens', str(agent.get('max_tokens', 'N/A')))
        ]
        
        # Batch fill table cells
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        doc.add_paragraph("")
    
    def _add_agent_performance_metrics_optimized(self, doc: Document, agent: Dict[str, Any]):
        """Add agent performance metrics as a table (optimized version)."""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Pre-calculate metrics
        success_rate = agent.get('success_rate', 0)
        success_rate_str = f"{success_rate*100:.1f}%" if success_rate else 'N/A'
        
        rows_data = [
            ('Total Queries', str(agent.get('total_queries', 0))),
            ('Average Response Time', f"{agent.get('avg_response_time_ms', 0):.0f}ms"),
            ('Success Rate', success_rate_str),
            ('Chain Type', agent.get('chain_type', 'basic'))
        ]
        
        # Batch fill table cells
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        doc.add_paragraph("")
    
    def _add_agent_basic_info(self, doc: Document, agent: Dict[str, Any]):
        """Add basic agent information as a table."""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Agent ID', str(agent.get('id', 'N/A'))),
            ('Model', agent.get('model_name', 'Unknown')),
            ('Created', agent.get('created_at', 'Unknown')[:10] if agent.get('created_at') else 'Unknown'),
            ('Status', 'Active' if agent.get('is_active', True) else 'Inactive'),
            ('Temperature', str(agent.get('temperature', 'N/A'))),
            ('Max Tokens', str(agent.get('max_tokens', 'N/A')))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_agent_performance_metrics(self, doc: Document, agent: Dict[str, Any]):
        """Add agent performance metrics as a table."""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Total Queries', str(agent.get('total_queries', 0))),
            ('Average Response Time', f"{agent.get('avg_response_time_ms', 0):.0f}ms"),
            ('Success Rate', f"{agent.get('success_rate', 0)*100:.1f}%" if agent.get('success_rate') else 'N/A'),
            ('Chain Type', agent.get('chain_type', 'basic'))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_chat_interaction(self, doc: Document, chat: Dict[str, Any], interaction_num: int):
        """Add a single chat interaction to the document."""
        doc.add_heading(f'Interaction {interaction_num}', level=2)
        
        # Timestamp
        timestamp = chat.get('timestamp', datetime.now())
        doc.add_paragraph(f"Time: {timestamp}")
        
        # Model used
        if chat.get('model_used'):
            doc.add_paragraph(f"Model: {chat['model_used']}")
        
        # Query type
        if chat.get('query_type'):
            doc.add_paragraph(f"Type: {chat['query_type']}")
        
        # User query
        doc.add_heading('User Query:', level=3)
        query_para = doc.add_paragraph(chat.get('user_query', 'No query recorded'))
        query_para.style = 'Quote'
        
        # Response
        doc.add_heading('AI Response:', level=3)
        response_para = doc.add_paragraph(chat.get('response', 'No response recorded'))
        response_para.style = 'Quote'
        
        # Response time
        if chat.get('response_time_ms'):
            doc.add_paragraph(f"Response Time: {chat['response_time_ms']/1000:.2f} seconds")
        
        doc.add_paragraph("")
    
    def _add_agent_responses(self, doc: Document, agent_responses: Dict[str, Any]):
        """Add agent responses to the document."""
        for i, (agent_name, response) in enumerate(agent_responses.items(), 1):
            doc.add_heading(f'{i}. {agent_name}', level=2)
            response_para = doc.add_paragraph(response)
            response_para.style = 'Quote'
            doc.add_paragraph("")
    
    def _add_debate_chain(self, doc: Document, debate_chain: List[Dict[str, Any]]):
        """Add debate chain results to the document."""
        for i, round_result in enumerate(debate_chain, 1):
            agent_name = round_result.get('agent_name', f'Agent {i}')
            doc.add_heading(f'Round {i}: {agent_name}', level=2)
            
            response = round_result.get('response', 'No response')
            response_para = doc.add_paragraph(response)
            response_para.style = 'Quote'
            
            if round_result.get('agent_id'):
                doc.add_paragraph(f"Agent ID: {round_result['agent_id']}")
            
            doc.add_paragraph("")
    
    def _add_compliance_details(self, doc: Document, details: Dict[str, Any]):
        """Add compliance check details to the document."""
        for idx, analysis in details.items():
            agent_name = analysis.get('agent_name', f'Agent {idx}')
            doc.add_heading(f'{agent_name}', level=2)
            
            reason = analysis.get('reason', analysis.get('raw_text', 'No analysis'))
            reason_para = doc.add_paragraph(reason)
            reason_para.style = 'Quote'
            doc.add_paragraph("")
    
    def _add_performance_metrics_table(self, doc: Document, metrics: Dict[str, Any]):
        """Add performance metrics as a table."""
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Total Time', f"{metrics.get('total_time_ms', 0):.1f}ms"),
            ('Retrieval Time', f"{metrics.get('retrieval_time_ms', 0):.1f}ms"),
            ('Generation Time', f"{metrics.get('generation_time_ms', 0):.1f}ms"),
            ('Documents Retrieved', str(metrics.get('documents_retrieved', 0))),
            ('Documents Used', str(metrics.get('documents_used', 0))),
            ('Relevance Score', f"{metrics.get('relevance_score', 0):.2f}"),
            ('Success', 'Yes' if metrics.get('success', False) else 'No')
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_quality_assessment_table(self, doc: Document, quality: Dict[str, Any]):
        """Add quality assessment as a table."""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Overall Quality', f"{quality.get('overall_quality', 0):.2f}"),
            ('Relevance', f"{quality.get('relevance_score', 0):.2f}"),
            ('Coherence', f"{quality.get('coherence_score', 0):.2f}"),
            ('Factual Accuracy', f"{quality.get('factual_accuracy', 0):.2f}"),
            ('Completeness', f"{quality.get('completeness_score', 0):.2f}"),
            ('Context Utilization', f"{quality.get('context_utilization', 0):.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_alignment_assessment_table(self, doc: Document, alignment: Dict[str, Any]):
        """Add alignment assessment as a table."""
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Intent Alignment', f"{alignment.get('intent_alignment_score', 0):.2f}"),
            ('Query Coverage', f"{alignment.get('query_coverage_score', 0):.2f}"),
            ('Instruction Adherence', f"{alignment.get('instruction_adherence_score', 0):.2f}"),
            ('Expected Answer Type', alignment.get('expected_answer_type', 'Unknown')),
            ('Actual Answer Type', alignment.get('answer_type_classification', 'Unknown')),
            ('Answer Type Match', 'Yes' if alignment.get('answer_type_match', False) else 'No'),
            ('Tone Consistency', f"{alignment.get('tone_consistency_score', 0):.2f}"),
            ('Scope Accuracy', f"{alignment.get('scope_accuracy_score', 0):.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        # Missing elements
        if alignment.get('missing_elements'):
            doc.add_paragraph("")
            doc.add_paragraph(f"Missing Elements: {', '.join(alignment['missing_elements'])}")
        
        # Extra elements
        if alignment.get('extra_elements'):
            doc.add_paragraph(f"Extra Elements: {', '.join(alignment['extra_elements'])}")
        
        doc.add_paragraph("")
    
    def _add_classification_metrics_table(self, doc: Document, classification: Dict[str, Any]):
        """Add classification metrics as a table."""
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Query Classification', classification.get('query_classification', 'Unknown')),
            ('Response Classification', classification.get('response_classification', 'Unknown')),
            ('Classification Confidence', f"{classification.get('classification_confidence', 0):.2f}"),
            ('Domain Relevance', classification.get('domain_relevance', 'Unknown')),
            ('Complexity Level', classification.get('complexity_level', 'Unknown')),
            ('Information Density', f"{classification.get('information_density', 0):.2f}"),
            ('Actionability Score', f"{classification.get('actionability_score', 0):.2f}"),
            ('Specificity Score', f"{classification.get('specificity_score', 0):.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _document_to_bytes(self, doc: Document) -> bytes:
        """Convert Document object to bytes."""
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    
    def cleanup_temp_files(self):
        """Clean up temporary files."""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                os.makedirs(self.temp_dir, exist_ok=True)
        except Exception as e:
            logger.warning(f"Failed to cleanup temp files: {str(e)}")
