import os
import re
import tempfile
import requests
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.shared import Inches
from fpdf import FPDF  
import streamlit as st
from io import BytesIO
from PIL import Image
import datetime
from io import BytesIO


# Set up environment variables and API endpoints
CHROMADB_API = os.getenv("CHROMA_URL", "http://localhost:8020")
LLM_API = os.getenv("LLM_API", "http://localhost:9020")
FASTAPI_API = os.getenv("FASTAPI_URL", "http://localhost:9020")
CHAT_ENDPOINT = f"{FASTAPI_API}/chat"
HISTORY_ENDPOINT = f"{FASTAPI_API}/chat-history"
HEALTH_ENDPOINT = f"{FASTAPI_API}/health"
OPEN_AI_API_KEY = os.getenv("OPEN_AI_API_KEY")


# Model configurations
model_key_map = {
    "GPT-4": "gpt-4",
    "GPT-3.5-Turbo": "gpt-3.5-turbo",
    "GPT-OSS": "gpt-oss",
    "LLaMA": "llama",  # single canonical LLaMA mapped to llama3.1:8b by backend
}

model_descriptions = {
    "gpt-4": "Most capable model for complex analysis",
    "gpt-3.5-turbo": "Cost-effective model for general tasks",
    "gpt-oss": "OpenAI OSS endpoints (uses OAI_OSS_MODEL)",
    "llama": "LLaMA via Ollama (defaults to llama3.1:8b)",
}

# Load Sentence Transformer model
embedding_model = SentenceTransformer('multi-qa-mpnet-base-dot-v1')


# Cache functions
@st.cache_data(ttl=1_200) #20 Minutes
def get_available_models_cached():
    return get_available_models()

def check_model_status(model_name):
    """Check if a specific model is loaded in Ollama"""
    try:
        response = requests.get(f"{FASTAPI_API}/health", timeout=5)
        if response.ok:
            health_data = response.json()
            models = health_data.get("models", {})
            return models.get(model_name, "unknown")
    except:
        return "unknown"

def upload_documents_to_chromadb(files, collection_name, openai_api_key=OPEN_AI_API_KEY):
    """Upload documents to ChromaDB using the existing endpoint"""
    try:
        # Prepare files for upload
        files_data = []
        for file in files:
            files_data.append(("files", (file.name, file.getvalue(), file.type)))
        
        # Prepare headers
        headers = {}
        if openai_api_key:
            headers["X-OpenAI-API-Key"] = openai_api_key
        
        # Prepare parameters
        params = {
            "collection_name": collection_name,
            "chunk_size": 1000,
            "chunk_overlap": 200,
            "store_images": True,
            "model_name": "enhanced",
            "debug_mode": False,
            "run_all_vision_models": True
        }
        
        # Make request to ChromaDB upload endpoint
        response = requests.post(
            f"{CHROMADB_API}/documents/upload-and-process",
            files=files_data,
            params=params,
            headers=headers,
            timeout=300
        )
        
        return response
        
    except Exception as e:
        raise Exception(f"Upload failed: {str(e)}")

def create_collection(collection_name):
    """Create a new ChromaDB collection"""
    try:
        response = requests.post(
            f"{CHROMADB_API}/collection/create",
            params={"collection_name": collection_name},
            timeout=30
        )
        return response
    except Exception as e:
        raise Exception(f"Collection creation failed: {str(e)}")
    
def get_chromadb_collections():
    """Get list of collections from ChromaDB"""
    try:
        response = requests.get(f"{CHROMADB_API}/collections", timeout=10)
        if response.ok:
            return response.json().get("collections", [])
        return []
    except:
        return []

def fetch_collections():
    try:
        response = requests.get(f"{CHROMADB_API}/collections", timeout=10)
        response.raise_for_status()
        return response.json().get("collections", [])
    except requests.exceptions.RequestException as e:
        print(f"Fetch collections failed: {e}")
        return []

def get_all_documents_in_collection(collection_name):
    """Get all documents in a collection with their metadata"""
    try:
        response = requests.get(
            f"{CHROMADB_API}/documents",
            params={"collection_name": collection_name},
            # timeout=300 # Increased timeout
        )
        if response.status_code == 200:
            data = response.json()
            
            # Group chunks by document_id to get unique documents
            documents = {}
            for i, doc_id in enumerate(data.get("ids", [])):
                metadata = data["metadatas"][i] if i < len(data.get("metadatas", [])) else {}
                document_id = metadata.get("document_id")
                document_name = metadata.get("document_name", "Unknown")
                
                if document_id and document_id not in documents:
                    documents[document_id] = {
                        "document_id": document_id,
                        "document_name": document_name,
                        "file_type": metadata.get("file_type", ""),
                        "total_chunks": metadata.get("total_chunks", 0),
                        "has_images": metadata.get("has_images", False),
                        "image_count": metadata.get("image_count", 0),
                        "processing_timestamp": metadata.get("timestamp", "")
                    }
            
            return list(documents.values())
        return []
    except Exception as e:
        st.error(f"Error fetching documents: {str(e)}")
        return []

def query_documents(collection_name, query_text, n_results=5):
    """Query documents using text search"""
    try:
        # Generate embedding for the query
        query_embedding = embedding_model.encode([query_text]).tolist()
        
        # Query ChromaDB
        response = requests.post(
            f"{CHROMADB_API}/documents/query",
            json={
                "collection_name": collection_name,
                "query_embeddings": query_embedding,
                "n_results": n_results,
                "include": ["documents", "metadatas", "distances"]
            },
            timeout=300
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            st.error(f"Query failed: {response.text}")
            return None
            
    except Exception as e:
        st.error(f"Error querying documents: {str(e)}")
        return None

def get_available_models():
    try:
        response = requests.get(f"{LLM_API}/health", timeout=10)
        if response.status_code == 200:
            health_data = response.json()
            if "models" in health_data:
                return [model for model, status in health_data["models"].items() if "available" in status.lower()]
    except Exception as e:
        print(f"Error fetching models: {e}")
    # Fallback list if health endpoint not available
    return ["gpt-4", "gpt-3.5-turbo", "gpt-oss", "llama"]


@st.cache_data(ttl=300)
def get_available_models_cached():
    return get_available_models()

def store_files_in_chromadb(files, collection_name, model_name="none", openai_api_key=None, 
                            chunk_size=1000, chunk_overlap=200, store_images=True, 
                            debug_mode=False, run_all_vision_models=True):
    """Store uploaded files in ChromaDB with enhanced parameters"""
    files_data = [('files', (file.name, file.getvalue(), file.type)) for file in files]
    
    params = {
        "collection_name": collection_name,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "store_images": store_images,
        "model_name": model_name,
        "debug_mode": debug_mode,
        "run_all_vision_models": run_all_vision_models  
    }
    
    # Prepare headers
    headers = {}
    if openai_api_key and model_name in ["gpt-4o-mini"]:
        headers["X-OpenAI-API-Key"] = openai_api_key
    
    try:
        response = requests.post(
            f"{CHROMADB_API}/documents/upload-and-process", 
            params=params, 
            files=files_data, 
            headers=headers, 
            timeout=600  
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to store documents: {response.status_code} - {response.text}")
        
        return response.json()
        
    except requests.exceptions.Timeout:
        raise Exception("Request timed out. Multi-model processing takes longer but provides comprehensive analysis.")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Could not connect to ChromaDB at {CHROMADB_API}")
    except Exception as e:
        raise Exception(f"Error storing files: {str(e)}")
    
def store_files_in_chromadb_selective(files, collection_name, model_name="none", openai_api_key=None, 
                                    chunk_size=1000, chunk_overlap=200, store_images=True, 
                                    debug_mode=False, selected_models=None):
    """Store uploaded files in ChromaDB with selective vision models"""
    files_data = [('files', (file.name, file.getvalue(), file.type)) for file in files]
    
    # Default to enhanced + basic if no models selected
    if not selected_models:
        selected_models = ["enhanced_local", "basic"]
    
    # Convert list to comma-separated string
    vision_models_str = ",".join(selected_models)
    
    params = {
        "collection_name": collection_name,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "store_images": store_images,
        "model_name": model_name,
        "debug_mode": debug_mode,
        "vision_models": vision_models_str  # Pass selected models
    }
    
    # Prepare headers
    headers = {}
    if openai_api_key and model_name in ["gpt-4o-mini"]:
        headers["X-OpenAI-API-Key"] = openai_api_key
    
    try:
        # Adjust timeout based on number of models selected
        base_timeout = 300
        model_timeout = len(selected_models) * 60  # 1 minute per model
        total_timeout = min(base_timeout + model_timeout, 900)  # Max 15 minutes
        
        response = requests.post(
            f"{CHROMADB_API}/documents/upload-and-process", 
            params=params, 
            files=files_data, 
            headers=headers, 
            timeout=total_timeout
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to store documents: {response.status_code} - {response.text}")
        
        return response.json()
        
    except requests.exceptions.Timeout:
        raise Exception(f"Request timed out. Processing {len(selected_models)} vision models takes significant time.")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Could not connect to ChromaDB at {CHROMADB_API}")
    except Exception as e:
        raise Exception(f"Error storing files: {str(e)}")


def query_documents_with_embedding(collection_name, query_text, n_results=5):
    """Query documents using text embedding"""
    try:
        # Generate embedding for the query
        query_embedding = embedding_model.encode([query_text]).tolist()
        
        # Query ChromaDB
        response = requests.post(
            f"{CHROMADB_API}/documents/query",
            json={
                "collection_name": collection_name,
                "query_embeddings": query_embedding,
                "n_results": n_results,
                "include": ["documents", "metadatas", "distances"]
            },
            timeout=60
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            raise Exception(f"Query failed: {response.text}")
            
    except Exception as e:
        raise Exception(f"Error querying documents: {str(e)}")

def get_all_documents_in_collection(collection_name):
    """Get all unique documents in a collection with their metadata"""
    try:
        response = requests.get(
            f"{CHROMADB_API}/documents",
            params={"collection_name": collection_name},
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to fetch documents: {response.text}")
            
        data = response.json()
        
        # Group chunks by document_id to get unique documents
        documents = {}
        for i, doc_id in enumerate(data.get("ids", [])):
            metadata = data["metadatas"][i] if i < len(data.get("metadatas", [])) else {}
            document_id = metadata.get("document_id")
            document_name = metadata.get("document_name", "Unknown")
            
            if document_id and document_id not in documents:
                documents[document_id] = {
                    "document_id": document_id,
                    "document_name": document_name,
                    "file_type": metadata.get("file_type", ""),
                    "total_chunks": metadata.get("total_chunks", 0),
                    # "has_images": metadata.get("has_images", False),
                    # "image_count": metadata.get("image_count", 0),
                    # "processing_timestamp": metadata.get("timestamp", ""),
                    # "openai_api_used": metadata.get("openai_api_used", False)
                }
        
        return list(documents.values())
        
    except Exception as e:
        raise Exception(f"Error fetching documents: {str(e)}")


def reconstruct_document_with_timeout(document_id, collection_name, timeout=300):
    """Reconstruct document with configurable timeout"""
    try:
        response = requests.get(
            f"{CHROMADB_API}/documents/reconstruct/{document_id}",
            params={"collection_name": collection_name},
            timeout=timeout
        )
        
        if response.status_code == 200:
            return response.json()
        elif response.status_code == 404:
            raise Exception("Document not found")
        else:
            raise Exception(f"Error: {response.text}")
            
    except requests.exceptions.Timeout:
        raise Exception("Request timed out. The document might be very large or the server is busy.")
    except Exception as e:
        raise Exception(f"Error reconstructing document: {str(e)}")

def get_image_from_chromadb(filename):
    """Get image from ChromaDB storage"""
    try:
        img_url = f"{CHROMADB_API}/images/{filename}"
        response = requests.get(img_url, timeout=30)
        
        if response.status_code == 200:
            return Image.open(BytesIO(response.content))
        else:
            return None
    except Exception as e:
        print(f"Error fetching image {filename}: {e}")
        return None

def render_reconstructed_document(result):
    """Renders the reconstructed text and inlines all images with st.image."""
    md = result["reconstructed_content"]

    # 1) Show the text first
    with st.expander("Reconstructed Document Text"):
        st.markdown(md, unsafe_allow_html=True)

    # 2) Then, for each image, re-use your existing logic
    if result.get("images"):
        with st.expander("Inline Images"):
            for i, img in enumerate(result["images"], 1):
                st.subheader(f"Image {i}: {img['filename']}")
                col1, col2 = st.columns([1, 2])

                # left: actual image
                with col1:
                    try:
                        resp = requests.get(f"{CHROMADB_API}/images/{img['filename']}")
                        resp.raise_for_status()
                        image = Image.open(BytesIO(resp.content))
                        st.image(image, caption=img['filename'] )
                    except:
                        st.write("Image preview not available")

                # right: metadata + description
                with col2:
                    st.write(f"**Storage Path:** `{img['storage_path']}`")
                    st.text_area("Description", img["description"], height=150, key=f"recon_desc_{i}")


def clean_text(text):
    """Remove null bytes and non-XML-safe characters"""
    if not text:
        return ""
    return ''.join(c for c in text if c == '\n' or c == '\t' or (32 <= ord(c) <= 126))

def export_to_docx(result):
    """Export document to DOCX with enhanced image handling"""
    return export_to_docx_with_markdown(result)

def export_to_pdf(result):
    """Export document to PDF with enhanced image handling"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    content = result["reconstructed_content"]
    images = result.get("images", [])
    image_map = {img["filename"]: img["description"] for img in images}

    # Split content by image markers
    parts = content.split("[IMAGE:")
    
    # Add first part (before any images)
    if parts[0].strip():
        pdf.multi_cell(0, 10, clean_text(parts[0].strip()))

    # Process each part that contains an image
    for part in parts[1:]:
        if "]" in part:
            filename, rest = part.split("]", 1)
            
            # Try to get and insert the image
            image_obj = get_image_from_chromadb(filename)
            if image_obj:
                try:
                    # Save image to temporary file
                    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    image_obj.save(temp_img.name)

                    # Add new page for image
                    pdf.add_page()
                    pdf.image(temp_img.name, w=150)
                    pdf.ln(5)
                    
                    # Add image description
                    if filename in image_map:
                        pdf.multi_cell(0, 10, clean_text(image_map[filename]))
                    
                    # Clean up temp file
                    os.unlink(temp_img.name)
                    
                except Exception as e:
                    pdf.multi_cell(0, 10, f"[Failed to insert image {filename}]: {e}")
            else:
                pdf.multi_cell(0, 10, f"[Image {filename} not found]")

            # Add remaining text after image
            if rest.strip():
                pdf.multi_cell(0, 10, clean_text(rest.strip()))
        else:
            # No image marker found, just add text
            if part.strip():
                pdf.multi_cell(0, 10, clean_text(part.strip()))

    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(temp_file.name)
    return temp_file.name

def check_chromadb_health():
    """Check ChromaDB health and return status info"""
    try:
        response = requests.get(f"{CHROMADB_API}/health", timeout=10)
        if response.status_code == 200:
            return True, response.json()
        return False, None
    except Exception as e:
        return False, str(e)
    
# Add these additional functions to your utils.py file

def create_file_like_object(content, filename):
    """Create a file-like object from text content for upload"""

    
    # Encode content as bytes
    content_bytes = content.encode('utf-8')
    
    # Create BytesIO object
    file_obj = BytesIO(content_bytes)
    file_obj.name = filename
    
    return file_obj

def markdown_to_docx(markdown_text, doc):
    """Convert markdown text to properly formatted DOCX content"""
    lines = markdown_text.split('\n')
    for line in lines:
        l = line.strip()
        if not l:
            # Add empty paragraph for spacing
            doc.add_paragraph("")
            continue
            
        # Handle headers with ## 
        if l.startswith("## "):
            doc.add_heading(l.replace("##", "").strip(), level=1)
        elif l.startswith("### "):
            doc.add_heading(l.replace("###", "").strip(), level=2)
        elif l.startswith("#### "):
            doc.add_heading(l.replace("####", "").strip(), level=3)
        # Handle lines that are fully bold as headings
        elif l.startswith("**") and l.endswith("**") and l.count("**") == 2:
            doc.add_heading(l.replace("*", "").strip(), level=2)
        # Handle bullet points
        elif l.startswith(("-", "*", "•")):
            doc.add_paragraph(l.lstrip("-*• ").strip(), style='List Bullet')
        # Handle numbered lists
        elif l[:2].isdigit() and l[2:4] in ('.', ') '):
            doc.add_paragraph(l.split('.', 1)[1].strip() if '.' in l else l.split(')', 1)[1].strip(), style='List Number')
        # Handle lines with mixed bold text
        elif "**" in l:
            parts = l.split("**")
            p = doc.add_paragraph()
            toggle = False
            for part in parts:
                if part:  # Only add non-empty parts
                    run = p.add_run(part)
                    if toggle:
                        run.bold = True
                toggle = not toggle
        # Handle horizontal rules
        elif l.startswith("---"):
            doc.add_paragraph("_" * 50)  # Add a line separator
        # Regular paragraph
        else:
            doc.add_paragraph(l)

def export_to_docx_with_markdown(result):
    """Enhanced export to DOCX with markdown formatting support"""
    try:
        from docx import Document
        import tempfile
        
        doc = Document()
        doc.add_heading(result["document_name"], 0)

        rich_content = result["reconstructed_content"]
        images = result.get("images", [])

        # Create image lookup for descriptions
        image_lookup = {img['filename']: img for img in images}

        # Replace image markers with placeholders for processing
        for img in images:
            marker = f"[IMAGE:{img['filename']}]"
            if marker in rich_content:
                rich_content = rich_content.replace(marker, f"@@IMAGE:{img['filename']}@@")

        # Split content by image markers
        chunks = re.split(r'(@@IMAGE:.+?@@)', rich_content)

        for chunk in chunks:
            img_match = re.match(r'@@IMAGE:(.+?)@@', chunk)
            if img_match:
                filename = img_match.group(1)
                
                # Try to get and insert the image
                image_obj = get_image_from_chromadb(filename)
                if image_obj:
                    try:
                        # Save image to temporary file for docx
                        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        image_obj.save(temp_img.name)
                        
                        # Add image to document
                        from docx.shared import Inches
                        doc.add_picture(temp_img.name, width=Inches(5.5))
                        
                        # Add image description if available
                        if filename in image_lookup:
                            desc = image_lookup[filename].get('description', '')
                            if desc:
                                doc.add_paragraph(clean_text(desc), style='Intense Quote')
                        
                        # Clean up temp file
                        os.unlink(temp_img.name)
                        
                    except Exception as e:
                        doc.add_paragraph(f"[Image '{filename}' could not be inserted: {e}]")
                else:
                    doc.add_paragraph(f"[Image '{filename}' not found]")
            else:
                # Process text chunk with markdown formatting
                cleaned_chunk = clean_text(chunk.strip())
                if cleaned_chunk:
                    # KEY FIX: Use the markdown_to_docx function for proper formatting
                    markdown_to_docx(cleaned_chunk, doc)

        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creating DOCX: {str(e)}")

def create_rule_document_docx(document_title, generated_content, metadata=None):
    """Create a DOCX file from generated rule content with markdown formatting"""
    try:
        # Create document
        doc = Document()
        doc.add_heading(document_title, 0)
        
        # Add metadata if provided
        if metadata:
            doc.add_heading('Document Information', 1)
            for key, value in metadata.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
            
            # Add separator
            doc.add_paragraph("")
        
        # Add main content heading
        doc.add_heading('Analysis Content', 1)
        
        # Process the generated content with markdown formatting
        # This is the key fix - actually convert the markdown
        markdown_to_docx(generated_content, doc)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creating DOCX: {str(e)}")

def create_combined_rule_documents_docx(generated_documents, collection_name):
    """Create a combined DOCX file from multiple generated rule documents"""
    try:
        # Create combined document
        doc = Document()
        doc.add_heading('Generated Rule Documents Collection', 0)
        
        # Add summary information
        doc.add_heading('Collection Summary', 1)
        doc.add_paragraph(f"Collection Name: {collection_name}")
        doc.add_paragraph(f"Total Documents: {len(generated_documents)}")
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add table of contents
        doc.add_heading('Table of Contents', 1)
        for i, doc_info in enumerate(generated_documents, 1):
            doc.add_paragraph(f"{i}. {doc_info['document_title']}")
        
        doc.add_page_break()
        
        # Add each document
        for i, doc_info in enumerate(generated_documents, 1):
            # Document header
            doc.add_heading(f"{i}. {doc_info['document_title']}", 1)
            
            # Document metadata
            metadata_info = [
                f"Source Document: {doc_info.get('source_document', 'Unknown')}",
                f"Generated by Agent: {doc_info.get('agent_name', 'Unknown')}",
                f"Rules Count: {doc_info.get('rules_count', 'Unknown')}",
                f"Content Length: {doc_info.get('content_length', 0):,} characters",
                f"Generated: {doc_info.get('generation_timestamp', 'Unknown')[:19]}"
            ]
            
            for info in metadata_info:
                doc.add_paragraph(info)
            
            doc.add_paragraph("")  # Empty line
            
            # Process document content with markdown formatting
            content = doc_info.get('generated_content', '') or doc_info.get('analysis_content', '')
            if content:
                # KEY FIX: Apply markdown conversion here
                markdown_to_docx(content, doc)
            
            # Add page break except for last document
            if i < len(generated_documents):
                doc.add_page_break()
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creating combined DOCX: {str(e)}")
    
def build_docx_bytes(generated_results):
    # 1) Create a new Document
    doc = Document()
    doc.add_heading('Generated Analysis Documents', level=0)

    # 2) Iterate and add each result
    for idx, result in enumerate(generated_results, 1):
        doc.add_heading(f"{idx}. {result['document_title']}", level=1)
        # insert metadata if you like
        doc.add_paragraph(f"Source: {result['source_document']}")
        doc.add_paragraph(f"Agent: {result['agent_name']}")
        doc.add_paragraph('')  # blank line
        # assume markdown_to_docx is already imported and converts your markdown
        markdown_to_docx(result['analysis_content'], doc)
        if idx < len(generated_results):
            doc.add_page_break()

    # 3) Save into a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
